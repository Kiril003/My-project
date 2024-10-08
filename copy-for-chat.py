import os
import sys
import logging
import shutil
import json
import time
import re
import queue
import wave
import random
import threading
import subprocess
import asyncio
import aiohttp
import importlib.util
import webbrowser
from datetime import datetime, timedelta
from enum import Enum
from functools import lru_cache
from typing import List, Dict, Any, Tuple, Optional
from concurrent.futures import ThreadPoolExecutor
import pickle


import keyboard
import mouse
import psutil
import pyautogui
import pygame
import requests
import torch
import numpy as np
import winshell
import geocoder
import sounddevice as sd
from scipy.signal import resample

import flet as ft
from flet import CupertinoContextMenu, CupertinoContextMenuAction

import speech_recognition as sr
from vosk import Model, KaldiRecognizer
import pyaudio

from deep_translator import GoogleTranslator
from num2words import num2words

from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume
from screen_brightness_control import set_brightness

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_UNDERLINE 
from docx.enum.section import WD_ORIENT, WD_SECTION, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt, RGBColor, Mm, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from PIL import Image, ImageFont, ImageDraw

from ctypes import cast, POINTER
from comtypes import CLSCTX_ALL
import comtypes

import wave
from fireworks.client import Fireworks

import seaborn as sns
import pandas as pd
import matplotlib.pyplot as plt


fireworks_client = Fireworks(api_key="fw_3ZNDV8Rj7jFi7mqCso2QztXL")
comtypes.CoInitialize()

COMMANDS_FILE = "commands.json"

CONFIG_FILE = "config.json"

torch.set_num_threads(4)
model = torch.package.PackageImporter('./model.pt').load_pickle("tts_models", "model")
model.to(torch.device('cpu'))
speaker='mykyta'
sample_rate=48000

SERVER_URL = "https://lionfish-unique-instantly.ngrok-free.app"
LOADED_CARDS_FILE = "loaded_cards.json"
CARDS_FUNCTIONALITY_FILE = "cards_functionality.json"

GLOBAL_USERNAME = None
GLOBAL_PASSWORD = None




def word_to_num(word):
    """Конвертує числівники в числа для української, англійської та російської мов."""
    numbers = {
        'нуль': 0, 'один': 1, 'два': 2, 'три': 3, 'чотири': 4, "п'ять": 5,
        'шість': 6, 'сім': 7, 'вісім': 8, "дев'ять": 9, 'десять': 10,
        'одинадцять': 11, 'дванадцять': 12, 'тринадцять': 13, 'чотирнадцять': 14,
        "п'ятнадцять": 15, 'шістнадцять': 16, 'сімнадцять': 17, 'вісімнадцять': 18,
        "дев'ятнадцять": 19, 'двадцять': 20, 'тридцять': 30, 'сорок': 40,
        "п'ятдесят": 50, 'шістдесят': 60, 'сімдесят': 70, 'вісімдесят': 80,
        "дев'яносто": 90, 'сто': 100,
        
        'ноль': 0, 'один': 1, 'два': 2, 'три': 3, 'четыре': 4, 'пять': 5,
        'шесть': 6, 'семь': 7, 'восемь': 8, 'девять': 9, 'десять': 10,
        'одиннадцать': 11, 'двенадцать': 12, 'тринадцать': 13, 'четырнадцать': 14,
        'пятнадцать': 15, 'шестнадцать': 16, 'семнадцать': 17, 'восемнадцать': 18,
        'девятнадцать': 19, 'двадцать': 20, 'тридцать': 30, 'сорок': 40,
        'пятьдесят': 50, 'шестьдесят': 60, 'семьдесят': 70, 'восемьдесят': 80,
        'девяносто': 90, 'сто': 100,

        'zero': 0, 'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
        'six': 6, 'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10,
        'eleven': 11, 'twelve': 12, 'thirteen': 13, 'fourteen': 14,
        'fifteen': 15, 'sixteen': 16, 'seventeen': 17, 'eighteen': 18,
        'nineteen': 19, 'twenty': 20, 'thirty': 30, 'forty': 40,
        'fifty': 50, 'sixty': 60, 'seventy': 70, 'eighty': 80,
        'ninety': 90, 'hundred': 100
    }
    
    words = word.lower().split()
    total = 0
    for i, w in enumerate(words):
        if w in numbers:
            if i < len(words) - 1 and words[i+1] in numbers:
                if numbers[w] >= 20 and numbers[words[i+1]] < 10:
                    total += numbers[w] + numbers[words[i+1]]
                    break
            total += numbers[w]
    return total if total else None

def extract_value(user_input, keywords):
    """Витягує і конвертує значення для вказаних ключових слів з тексту."""
    pattern = '|'.join(keywords)
    match = re.search(fr'({pattern})\s+(на|to|on)\s+([\w\s]+)', user_input, re.IGNORECASE)
    if match:
        value = match.group(3).strip().replace('\n', ' ')
        value = re.sub(r'\s+та\s+|\s+and\s+|\s+и\s+', ' ', value)
        try:
            return int(value)
        except ValueError:
            return word_to_num(value)
    return None

def extract_value(user_input, keywords):
    """Витягує і конвертує значення для вказаних ключових слів з тексту."""
    pattern = '|'.join(keywords)
    match = re.search(fr'({pattern})\s+(на|to|on)\s+([\w\s]+)', user_input, re.IGNORECASE)
    if match:
        value = match.group(3).strip().replace('\n', ' ')
        value = re.sub(r'\s+та\s+|\s+and\s+|\s+и\s+', ' ', value)
        try:
            return int(value)
        except ValueError:
            return word_to_num(value)
    return None

def change_system_volume(value):
    """Змінює системну гучність на задане значення (0-100)."""
    devices = AudioUtilities.GetSpeakers()
    interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
    volume = interface.QueryInterface(IAudioEndpointVolume)
    volume.SetMasterVolumeLevelScalar(value / 100.0, None)

def change_screen_brightness(value):
    """Змінює яскравість екрану на задане значення (0-100)."""
    set_brightness(value)

def process_input(user_input):
    """Обробляє ввід користувача і змінює гучність або яскравість."""
    volume_keywords = ['гучність', 'volume', 'громкость']
    brightness_keywords = ['яскравість', 'brightness', 'яркость']

    volume = extract_value(user_input, volume_keywords)
    brightness = extract_value(user_input, brightness_keywords)

    if volume is not None:
        if 0 <= volume <= 100:
            change_system_volume(volume)
            print(f"Гучність змінено на {volume}%")
        else:
            print("Некоректне значення гучності. Введіть число від 0 до 100 або число словами.")

    if brightness is not None:
        if 0 <= brightness <= 100:
            change_screen_brightness(brightness)
            print(f"Яскравість змінено на {brightness}%")
        else:
            print("Некоректне значення яскравості. Введіть число від 0 до 100 або число словами.")

class WeatherInfo:
    BASE_URL = "http://api.openweathermap.org/data/2.5/weather"
    
    DAYS = ['понеділок', 'вівторок', 'середа', 'четвер', 'п\'ятниця', 'субота', 'неділя']
    
    TIME_INDICATORS = {
        'сьогодні': 0,
        'завтра': 1,
        'післязавтра': 2,
        'наступного тижня': 7
    }
    
    WEATHER_RECOMMENDATIONS = {
        'clear': [
            "Чудовий день для прогулянки!",
            "Насолоджуйтесь сонячною погодою!",
            "Ідеальний час для пікніка!",
            "Чудова погода для велопрогулянки!",
            "Насолоджуйтесь днем на природі!"
        ],
        'clouds': [
            "Хороший день для фотосесії!",
            "Приємна погода для прогулянки.",
            "Можливо, варто взяти легку куртку.",
            "Чудова нагода для читання на свіжому повітрі!",
            "Хмарно, але все одно приємно!"
        ],
        'rain': [
            "Не забудьте парасольку!",
            "Гарний день, щоб залишитися вдома з книгою.",
            "Дощовик стане в нагоді!",
            "Погода для прогулянки під дощем!",
            "Час для кави під дощем!"
        ],
        'snow': [
            "Час ліпити сніговика!",
            "Одягніться тепліше.",
            "Будьте обережні на дорогах!",
            "Ідеальний час для зимових розваг!",
            "Зима чарівна, насолоджуйтесь!"
        ],
        'extreme': [
            "Краще залишитися вдома.",
            "Будьте дуже обережні надворі!",
            "Слідкуйте за повідомленнями від служб надзвичайних ситуацій.",
            "Небезпека! Залишайтеся вдома!",
            "Уникайте подорожей у цю погоду."
        ]
    }

    def __init__(self, api_key: str):
        self.api_key = api_key

    def parse_command(self, command: str) -> Tuple[Optional[str], Optional[datetime], Optional[str]]:
        command = command.lower()
        words = command.split()
        city = None
        date = datetime.now()
        day = None

        # Видаляємо слово "погода" та прийменники
        words = [word for word in words if word not in ['погода', 'в', 'у', 'на', 'місті', 'місто']]

        # Шукаємо часові індикатори
        time_indicators = list(self.TIME_INDICATORS.keys()) + self.DAYS
        time_index = next((i for i, word in enumerate(words) if word in time_indicators or any(indicator in word for indicator in time_indicators)), len(words))

        # Визначаємо місто
        city = " ".join(words[:time_index])

        # Визначаємо дату або день
        for day_name in self.DAYS:
            if day_name in command:
                day = day_name
                date = self.get_next_day_date(day_name)
                break

        if not day:
            for time_indicator, days_ahead in self.TIME_INDICATORS.items():
                if time_indicator in command:
                    date = datetime.now() + timedelta(days=days_ahead)
                    break

        # Нормалізуємо назву міста (прибираємо відмінки)
        if city.endswith('і'):
            city = city[:-1] + 'а'  # Змінюємо "Києві" на "Київ", "Одесі" на "Одеса" і т.д.
        elif city.endswith('ку'):
            city = city[:-2] + 'к'  # Змінюємо "Луцьку" на "Луцьк"

        return city.strip(), date, day

    @staticmethod
    def get_next_day_date(day_name: str) -> datetime:
        today = datetime.now()
        day_index = WeatherSystem.DAYS.index(day_name)
        days_ahead = (day_index - today.weekday() + 7) % 7
        if days_ahead == 0:
            days_ahead = 7
        return today + timedelta(days=days_ahead)

    @lru_cache(maxsize=128)
    def get_weather(self, city: str) -> Tuple[Optional[float], Optional[str], Optional[str]]:
        params = {
            'q': city,
            'appid': self.api_key,
            'units': 'metric',
            'lang': 'ua'
        }
        try:
            response = requests.get(self.BASE_URL, params=params, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            temp = data['main']['temp']
            description = data['weather'][0]['description']
            main_weather = data['weather'][0]['main'].lower()
            return temp, description, main_weather
        
        except requests.RequestException as e:
            if isinstance(e, requests.HTTPError) and e.response.status_code == 404:
                return None, f"Місто '{city}' не знайдено. Перевірте правильність написання.", None
            return None, f"Помилка при запиті до API: {e}", None
        except (ValueError, KeyError) as e:
            return None, f"Помилка при обробці даних: {e}", None
        except Exception as e:
            return None, f"Несподівана помилка: {e}", None
    
    def get_recommendation(self, main_weather: str) -> str:
        weather_type = 'extreme'
        for key in self.WEATHER_RECOMMENDATIONS:
            if main_weather.startswith(key):
                weather_type = key
                break
        return random.choice(self.WEATHER_RECOMMENDATIONS[weather_type])

    def process_weather_command(self, command: str) -> Tuple[str, str]:
        city, date, day = self.parse_command(command)
        
        if not city:
            return "Вибачте, не вдалося розпізнати місто. Спробуйте ще раз, вказавши місто явно, наприклад 'погода в місті Київ'.", ""

        temp, description, main_weather = self.get_weather(city)
        
        if temp is None:
            return description, ""

        recommendation = self.get_recommendation(main_weather)
        
        response1 = self.format_weather_response(city, date, day, temp, description, recommendation, format_type=1)
        response2 = self.format_weather_response(city, date, day, temp, description, recommendation, format_type=2)
        
        return response1, response2

    def format_weather_response(self, city: str, date: datetime, day: Optional[str], temp: float, description: str, recommendation: str, format_type: int) -> str:
        today = datetime.now()
        if date.date() == today.date():
            date_str = "сьогодні"
        elif date.date() == (today + timedelta(days=1)).date():
            date_str = "завтра"
        elif date.date() == (today + timedelta(days=2)).date():
            date_str = "післязавтра"
        elif day:
            date_str = f"у {day}"
        else:
            date_str = f"на {self.DAYS[date.weekday()]}"

        if format_type == 1:
            temp_str = f"{temp:.1f}°C"
        else:
            temp_int = int(round(temp))
            temp_str = f"{temp_int} {'градус' if temp_int % 10 == 1 and temp_int % 100 != 11 else 'градуси' if 2 <= temp_int % 10 <= 4 and (temp_int % 100 < 10 or temp_int % 100 >= 20) else 'градусів'}"

        return f"У місті {city.capitalize()} {date_str} очікується температура {temp_str}, {description}. {recommendation}"

weather_info = WeatherInfo(api_key='78de1db61ffa6efd32239911ca57f068')

class Card:
    def __init__(self, id, image, description, functionality, content=None):
        self.id = id
        self.image = image
        self.description = description
        self.functionality = functionality
        self.content = content

    def to_dict(self):
        return {
            "id": self.id,
            "image": self.image,
            "description": self.description,
            "functionality": self.functionality,
            "content": self.content
        }

async def fetch_cards():
    async with aiohttp.ClientSession() as session:
        async with session.get(f"{SERVER_URL}/cards") as response:
            return await response.json()

async def fetch_card_functionality(card_id):
    async with aiohttp.ClientSession() as session:
        async with session.get(f"{SERVER_URL}/functionality/{card_id}") as response:
            return await response.json()

def save_loaded_cards(cards):
    with open(LOADED_CARDS_FILE, 'w') as f:
        json.dump([card.to_dict() for card in cards], f)
    print(f"Saved {len(cards)} cards to {LOADED_CARDS_FILE}")  # Додайте цей рядок для налагодження

def load_saved_cards():
    if os.path.exists(LOADED_CARDS_FILE):
        with open(LOADED_CARDS_FILE, 'r') as f:
            loaded_cards = [Card(**card_data) for card_data in json.load(f)]
        print(f"Loaded {len(loaded_cards)} cards from {LOADED_CARDS_FILE}")  # Додайте цей рядок для налагодження
        return loaded_cards
    return []

def load_config():
    default_config = {
        "language": "uk-UA",
        "use_vosk": False,
        "preferred_recognition": "google",
        "vosk_models": {
            "uk-UA": "",
            "en-US": "",
            "ru-RU": ""
        },
        "enable_tts": True,
        "use_ai": False,
        "ai_model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "flux_model": "black-forest-labs/FLUX.1-dev",
        "text_generation_model": "custom"
    }

    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r') as f:
            config = json.load(f)
            
            # Оновлюємо конфігурацію, додаючи відсутні ключі
            for key, value in default_config.items():
                if key not in config:
                    config[key] = value

            if "preferred_recognition" not in config:
                config["preferred_recognition"] = "vosk" if config.get("use_vosk", False) else "google"

            if "use_ai" not in config:
                config["use_ai"] = False

            if "ai_model" not in config:
                config["ai_model"] = "mistralai/Mixtral-8x7B-Instruct-v0.1"

        return config
    return default_config

def save_config(config):
    with open(CONFIG_FILE, 'w') as f:
        json.dump(config, f, indent=4)



def preprocess_text(text):
        # Словник для заміни англійських слів на транскрипцію
        eng_dict = {
            '+': 'плюс', '×':'помножити на', '÷':'ділити на', '=':'дорівнює', '%':'відсоток',
            'was': 'воз', 'were': 'вер', 'be': 'бі', 'been': 'бін', 'being': 'бійнг',
            'have': 'хев', 'has': 'хез', 'had': 'хед', 'do': 'ду', 'does': 'даз',
            'did': 'дід', 'will': 'віл', 'would': 'вуд', 'shall': 'шел', 'should': 'шуд',
            'can': 'кен', 'could': 'куд', 'may': 'мей', 'might': 'майт',
            'must': 'маст', 'ought': 'от', 'and': 'енд', 'or': 'ор', 'but': 'бат',
            'because': 'бікоз', 'if': 'іф', 'although': 'олзо', 'unless': 'анлес',
            'since': 'сінс', 'for': 'фор', 'yet': 'йет', 'so': 'соу',
            'as': 'ез', 'nor': 'нор', 'while': 'вайл', 'where': 'вер',
            'when': 'вен', 'how': 'хау', 'what': 'вот', 'who': 'ху',
            'whom': 'хум', 'whose': 'хуз', 'which': 'віч', 'why': 'вай',
            'hello': 'хелоу', 'hi': 'хай', 'bye': 'бай', 'goodbye': 'гудбай',
            'yes': 'єс', 'no': 'ноу', 'okay': 'окей', 'ok': 'окей',
            'please': 'пліз', 'thank': 'сенк', 'you': 'ю', 'he': 'хі',
            'she': 'ші', 'it': 'іт', 'we': 'ві', 'they': 'зей',
            'me': 'мі', 'him': 'хім', 'her': 'хер', 'us': 'ас',
            'them': 'зем', 'my': 'май', 'your': 'йор', 'his': 'хіз',
            'its': 'ітс', 'our': 'ауер', 'their': 'зеір', 'this': 'зіс',
            'that': 'зет', 'these': 'зіз', 'those': 'зоуз', 'here': 'хіер',
            'there': 'зеер', 'now': 'нау', 'then': 'зен', 'today': 'тудей',
            'tomorrow': 'тумороу', 'yesterday': 'єстедей', 'always': 'олвейз',
            'never': 'невер', 'sometimes': 'самтаймз', 'often': 'офтен',
            'usually': 'южуалі', 'rarely': 'рерлі', 'seldom': 'селдом',
            'good': 'гуд', 'bad': 'бед', 'big': 'біг', 'small': 'смол',
            'large': 'лардж', 'little': 'літл', 'long': 'лонг', 'short': 'шорт',
            'high': 'хай', 'low': 'лоу', 'hot': 'хот', 'cold': 'колд',
            'warm': 'ворм', 'cool': 'кул', 'new': 'нью', 'old': 'олд',
            'young': 'янг', 'beautiful': 'бютіфул', 'ugly': 'аглі', 'happy': 'хепі',
            'sad': 'сед', 'angry': 'енгрі', 'tired': 'тайерд', 'hungry': 'хангрі',
            'thirsty': 'серсті', 'rich': 'річ', 'poor': 'пур', 'busy': 'бізі',
            'free': 'фрі', 'strong': 'стронг', 'weak': 'вік', 'true': 'тру',
            'false': 'фолс', 'right': 'райт', 'wrong': 'ронг', 'easy': 'ізі',
            'difficult': 'діфікалт', 'fast': 'фаст', 'slow': 'слоу', 'early': 'ерлі',
            'late': 'лейт', 'loud': 'лауд', 'quiet': 'квайєт', 'hard': 'хард',
            'soft': 'софт', 'clean': 'клін', 'dirty': 'дьорті', 'dry': 'драй',
            'wet': 'вет', 'full': 'фул', 'empty': 'емпті', 'light': 'лайт',
            'heavy': 'хеві', 'open': 'оупен', 'closed': 'клоузд', 'near': 'ніар',
            'far': 'фар', 'deep': 'діп', 'shallow': 'шелоу', 'wide': 'вайд',
            'narrow': 'нероу', 'thick': 'сік', 'thin': 'сін', 'straight': 'стрейт',
            'curved': 'кьорвд', 'flat': 'флет', 'round': 'раунд', 'sharp': 'шарп',
            'dull': 'дал', 'smooth': 'смуз', 'rough': 'раф', 'tight': 'тайт',
            'loose': 'луз', 'correct': 'корект', 'incorrect': 'інкорект',
            'important': 'імпортент', 'unimportant': 'анімпортент', 'necessary': 'несесері',
            'unnecessary': 'аннесесері', 'possible': 'посібл', 'impossible': 'імпосібл',
            'safe': 'сейф', 'dangerous': 'дейнджерос', 'legal': 'лігал', 
            'illegal': 'іллігал', 'public': 'паблік', 'private': 'прайвіт',
            'expensive': 'експенсів', 'cheap': 'чіп', 'popular': "поп'юлар",
            'unpopular': "анпоп'юлар", 'interesting': 'інтерестінг', 'boring': 'борінг',
            'polite': 'полайт', 'rude': 'руд', 'kind': 'кайнд', 'cruel': 'круел',
            'brave': 'брейв', 'cowardly': 'кауердлі', 'clever': 'клевер', 'stupid': 'стьюпід',
            'wise': 'вайз', 'foolish': 'фуліш', 'honest': 'онест', 'dishonest': 'дісонест',
            'generous': 'дженерос', 'greedy': 'гріді', 'patient': 'пейшент',
            'impatient': 'імпейшент', 'calm': 'калм', 'nervous': 'нервос',
            'serious': 'сіріос', 'funny': 'фанні', 'formal': 'формал', 'informal': 'інформал',
            'modern': 'модерн', 'traditional': 'традішнл', 'simple': 'сімпл',
            'complicated': 'комплікейтед', 'natural': 'нечурал', 'artificial': 'артіфішл',
            'physical': 'фізікал', 'mental': 'ментал', 'healthy': 'хелсі', 'sick': 'сік',
            'alive': 'елайв', 'dead': 'дед', 'awake': 'евейк', 'asleep': 'еслеп',
            'wet': 'вет', 'dry': 'драй', 'sweet': 'світ', 'sour': 'сауер',
            'bitter': 'біттер', 'salty': 'солті', 'spicy': 'спайсі', 'tasty': 'тейсті',
            'delicious': 'делішес', 'disgusting': 'дісгастінг',
            'a': 'ей', 'an': 'ен', 'the': 'зе', 'is': 'із', 'are': 'ар',
            'a': 'ей', 'b': 'б', 'c': 'с', 'd': 'д', 'e': 'і',
            'f': 'ф', 'g': 'г', 'h': 'х', 'i': 'ай', 'j': 'дж',
            'k': 'к', 'l': 'л', 'm': 'м', 'n': 'н', 'o': 'оу',
            'p': 'п', 'q': 'кью', 'r': 'ар', 's': 'ес', 't': 'ті',
            'u': 'ю', 'v': 'ві', 'w': 'в', 'x': 'екс', 'y': 'ай', 'z': 'з',
        }
        
        # Словник для заміни чисел місяця
        month_dict = {
            "січня": "першого", "лютого": "другого", "березня": "третього",
            "квітня": "четвертого", "травня": "п'ятого", "червня": "шостого",
            "липня": "сьомого", "серпня": "восьмого", "вересня": "дев'ятого",
            "жовтня": "десятого", "листопада": "одинадцятого", "грудня": "дванадцятого"
        }

        # Функція для заміни чисел словами
        def number_to_words(match):
            number = int(match.group())
            if 0 <= number <= 1000000:
                return num2words(number, lang='uk')
            return match.group()

        # Функція для обробки дат
        def process_date(match):
            day, month, year = match.groups()
            day_word = num2words(int(day), lang='uk')
            month_word = month_dict.get(month, month)
            year_word = ' '.join([num2words(int(d), lang='uk') for d in year])
            return f"{day_word} {month_word} {year_word} року"

        # Функція для обробки часу
        def process_time(match):
            hours, minutes = match.groups()
            hours_word = num2words(int(hours), lang='uk')
            minutes_word = num2words(int(minutes), lang='uk')
            return f"{hours_word} годин{'а' if hours == '1' else 'и'} {minutes_word} хвилин"

        # Функція для обробки дат у форматі YYYY-MM-DD
        def process_iso_date(match):
            year, month, day = match.groups()
            year_word = ' '.join([num2words(int(d), lang='uk') for d in year])
            month_word = num2words(int(month), lang='uk')
            day_word = num2words(int(day), lang='uk')
            return f"{day_word} {month_word} {year_word} року"

        # Заміна англійських слів
        for eng, ukr in eng_dict.items():
            text = re.sub(r'\b' + re.escape(eng) + r'\b', ukr, text, flags=re.IGNORECASE)

        # Заміна чисел словами
        text = re.sub(r'\b\d+\b', number_to_words, text)

        # Обробка дат
        text = re.sub(r'(\d{1,2})\s+(січня|лютого|березня|квітня|травня|червня|липня|серпня|вересня|жовтня|листопада|грудня)\s+(\d{4})', process_date, text)

        # Обробка часу
        text = re.sub(r'(\d{1,2}):(\d{2})', process_time, text)

        # Обробка дат у форматі YYYY-MM-DD
        text = re.sub(r'(\d{4})-(\d{2})-(\d{2})', process_iso_date, text)

        return text



def setup_groq_client():
    try:
        return Groq(api_key="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")
    except Exception as e:
        logging.error(f"Помилка під час налаштування клієнта Groq: {str(e)}")
        return None
    

# def speaks(text):
#     def _speak_in_thread(text):
#         config = load_config()
#         if config.get("enable_tts", True):
#             try:
#                 # Попередня обробка тексту
#                 processed_text = preprocess_text(text)
                
#                 audio = model.apply_tts(text=processed_text,
#                                         speaker=speaker,
#                                         sample_rate=sample_rate,
#                                         put_accent=True,
#                                         put_yo=True)
#                 audio = audio.view(-1)
#                 audio = torch.nn.functional.interpolate(
#                     audio.unsqueeze(0).unsqueeze(0),
#                     scale_factor=1,
#                     mode='linear',
#                     align_corners=False
#                 ).squeeze()
#                 speed_factor = 1
#                 num_samples = int(len(audio) / speed_factor)
#                 audio_resampled = resample(audio.numpy(), num_samples)
#                 sd.play(audio_resampled, sample_rate)
#                 sd.wait()
#                 sd.stop()
#             except Exception as e:
#                 print(f"Помилка озвучування: {str(e)}")

#     # Запуск озвучування в окремому потоці
#     threading.Thread(target=_speak_in_thread, args=(text,)).start()

# speaks("Вітаю")
# speaks("Як ти справишся?")
# speaks("Оламіесь")

class WeatherSystem:
    BASE_URL = "http://api.openweathermap.org/data/2.5/weather"
    
    def __init__(self, api_key):
        self.api_key = api_key
        self.weather_data = None
        self.last_update = None

    def get_local_weather(self):
        try:
            g = geocoder.ip('me')
            lat, lon = g.latlng
            
            params = {
                'lat': lat,
                'lon': lon,
                'appid': self.api_key,
                'units': 'metric',
                'lang': 'en'
            }
            
            response = requests.get(self.BASE_URL, params=params)
            response.raise_for_status()
            data = response.json()
            
            self.weather_data = {
                'city': data['name'],
                'country': data['sys']['country'],
                'temp': data['main']['temp'],
                'feels_like': data['main']['feels_like'],
                'humidity': data['main']['humidity'],
                'description': data['weather'][0]['description'],
                'wind_speed': data['wind']['speed']
            }
            self.last_update = datetime.now()
            
        except Exception as e:
            print(f"Error fetching weather data: {e}")

    def get_weather_string(self):
        if not self.weather_data:
            return "Weather data unavailable."
        
        return f"Current weather in {self.weather_data['city']}, {self.weather_data['country']}: " \
               f"Temperature: {self.weather_data['temp']}°C, " \
               f"Feels like: {self.weather_data['feels_like']}°C, " \
               f"Humidity: {self.weather_data['humidity']}%, " \
               f"Conditions: {self.weather_data['description']}, " \
               f"Wind speed: {self.weather_data['wind_speed']} m/s. " \
               f"Last updated: {self.last_update.strftime('%Y-%m-%d %H:%M:%S')}"

    def update_weather(self):
        self.get_local_weather()
        threading.Timer(1800, self.update_weather).start()  # 1800 seconds = 30 minutes

weather_system = WeatherSystem(api_key='78de1db61ffa6efd32239911ca57f068')
weather_info = weather_system.get_weather_string()
groq_client = setup_groq_client()

def generate_system_prompt():
    try:
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        return f"""You are an assistant named ART, designed to provide concise, accurate, and clear answers to user inquiries. To maintain a pleasant atmosphere, include emojis where appropriate. Responses should be brief unless the user requests detailed explanations, stories, or essays. 
- Default browser: Chrome
- Default operating system: Windows
- For image generation requests, respond with 'generate a photo ...(only in english)'
- For volume or brightness adjustments, execute using the 'subprocess' library.
- When asked to run commands on a PC, provide Python code wrapped between [commands] and [/commands] tags, ensuring delays between certain commands. Example:
[commands]
os.system(f'start {{command_name}}')
os.system('start ms-settings:..')
[/commands]

Current date and time: {current_time}
Current weather information: {weather_info}

- For creating documents or presentations, use Microsoft Word (not PowerPoint) and the 'docx' library. Ensure text is formatted in Ukrainian, while file names remain in English. Example:
[commands]
# Для покращення візуального представлення даних у документі, використовуй бібліотеки, такі як matplotlib або seaborn, для створення складних графіків (графіки обов'язкові у презентаціях де є певні дані). Нижче представлен приклад графіків опираючись а які ти зможеш створити необхідний або покращений:
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# Приклад створення графіка
data = sns.load_dataset('tips')
plt.figure(figsize=(6, 4))
sns.boxplot(x='day', y='total_bill', hue='smoker', data=data, palette='coolwarm')
plt.title('Total Bill Distribution by Day and Smoking Status')
plt.tight_layout()
plt.savefig('chart.png')
plt.close()

# Створення документа і вставка графіка
document = Document()
document.add_heading("Звіт з графіком", 0)
document.add_picture('chart.png', width=Inches(5))

file_name = 'report_with_chart.docx'

if os.path.exists(file_name):
    os.remove(file_name)

document.save(file_name)
os.startfile(file_name)

# приклад графіку 2:
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Створення даних та складного графіка
df = pd.DataFrame({{
    'Month': ['Jan', 'Feb', 'Mar', 'Apr', 'May'],
    'Sales': [1500, 1600, 1700, 1800, 1900],
    'Profit': [200, 220, 210, 230, 250]
}})

fig, ax1 = plt.subplots()

# Побудова лінійного графіка для прибутку
ax1.plot(df['Month'], df['Profit'], color='green', marker='o')
ax1.set_ylabel('Profit', color='green')

# Побудова гістограми для продажів
ax2 = ax1.twinx()
ax2.bar(df['Month'], df['Sales'], alpha=0.6, color='blue')
ax2.set_ylabel('Sales', color='blue')

plt.title('Sales and Profit Over Months')
plt.tight_layout()

# Збереження графіка
plt.savefig('complex_table_chart.png')
plt.close()

# Створення документу Word і вставка графіка
doc = Document()
doc.add_heading('Sales and Profit Report', 0)
doc.add_picture('complex_table_chart.png', width=Inches(5))
doc.save('complex_table_report.docx')
[/commands]

- When creating presentations, ensure they are well-structured and contain more than just a couple of sentences. A proper presentation should include:
  1. **Title Slide**: The title of the presentation, the author's name, and the date.
  2. **Introduction Slide**: An introduction to the topic, outlining the key points to be covered.
  3. **Content Slides**: Several very big slides (at least 5-7) that delve into the main content, each focusing on different aspects of the topic.
  4. **Visuals**: Use relevant images, diagrams, or charts to support the content.
  5. **Conclusion Slide**: A summary of the key points and any conclusions drawn.
Avoid embedding photos directly into Word files; instead, generate them separately and add them as needed.

Note: When detailed factual information is requested, use Brave Search to find and confirm the information before responding.
-The answer must be Ukrainian only!
<|eot_id|><|start_header_id|>user<|end_header_id|>
"""
    except Exception as e:
        logging.error(f"Error generating system prompt: {str(e)}")
        return ""


def remove_emoji(text):
    return re.sub(r'[^\w\s]', '', text)

def process_response(response):
    try:
        processed_response = response
        while True:
            start_idx = processed_response.find("[commands]")
            end_idx = processed_response.find("[/commands]")
            if start_idx == -1 or end_idx == -1:
                break

            command_string = processed_response[start_idx + len("[commands]"):end_idx].strip()
            try:
                exec(command_string)
            except Exception as e:
                logging.error(f"Сталася помилка під час виконання команди: {str(e)}")

            processed_response = processed_response[:start_idx] + processed_response[end_idx + len("[/commands]"):]
        return processed_response.strip()
    except Exception as e:
        logging.error(f"Помилка при обробці відповіді: {str(e)}")
        return response

def extract_emoji_and_text(text):
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"  
        u"\U0001F300-\U0001F5FF"  
        u"\U0001F680-\U0001F6FF"  
        u"\U0001F1E0-\U0001F1FF"  
        u"\U00002702-\U000027B0"
        u"\U000024C2-\U0001F251"
        "]+", flags=re.UNICODE)
    parts = []
    last_end = 0
    for match in emoji_pattern.finditer(text):
        if match.start() > last_end:
            parts.append(('text', text[last_end:match.start()]))
        parts.append(('emoji', match.group()))
        last_end = match.end()
    if last_end < len(text):
        parts.append(('text', text[last_end:]))
    return parts

from functools import lru_cache

@lru_cache(maxsize=100)
def cached_translate(text, source, target):
    return GoogleTranslator(source=source, target=target).translate(text)

def translate_text(text, source, target):
    try:
        parts = extract_emoji_and_text(text)
        translated_parts = []
        
        for part_type, part_text in parts:
            if part_type == 'emoji':
                translated_parts.append(part_text)
            else:
                translated_text = cached_translate(part_text, source, target)
                translated_parts.append(translated_text)
        
        return ''.join(translated_parts)
    except Exception as e:
        logging.error(f"Помилка при перекладі з {source} на {target}: {str(e)}")
        return text

def handle_request(messages, history, command_processor):
    try:
        if fireworks_client is None:
            raise Exception("Fireworks клієнт не був ініціалізований")

        translated_messages = []
        
        # Додаємо історію без перекладу
        for h in history[-2:]:
            translated_messages.append({"role": "user" if h.startswith("user") else "assistant", "content": h})
        
        # Додаємо тільки нове повідомлення без перекладу
        for m in messages:
            translated_messages.append(m)
        
        chat_completion = fireworks_client.chat.completions.create(
            messages=translated_messages,
            model="accounts/fireworks/models/llama-v3p1-405b-instruct",
            max_tokens=3500
        )
        
        if not hasattr(chat_completion, 'choices') or not chat_completion.choices:
            raise Exception("Unexpected response format from Fireworks API")
        
        response = chat_completion.choices[0].message.content
        logging.info(f"Помічник (від моделі): {response}")

        # Обробляємо відповідь перед виведенням
        processed_response = process_response(response)
        
        # Виводимо оброблену відповідь у термінал
        print(f"Оброблена відповідь моделі: {processed_response}")

        # Перевіряємо, чи містить відповідь запит на створення зображення
        image_prompt = None
        if "create a photo" in processed_response.lower():
            image_prompt = processed_response.lower().split("create a photo", 1)[1].strip()
        elif "generate a photo" in processed_response.lower():
            image_prompt = processed_response.lower().split("generate a photo", 1)[1].strip()

        if image_prompt:
            processed_response = f"Зображення згенероване за запитом: {image_prompt}"
            # Викликаємо метод generate_image з класу CommandProcessor
            command_processor.generate_image(f"згенеруй зображення {image_prompt}")

        return processed_response.replace("Помічник:".lower(), "", 1).strip()
    
    except Exception as e:
        logging.error(f"Помилка при обробці запиту: {str(e)}")
        if "503" in str(e):
            return "Сервіс тимчасово недоступний. Будь ласка, спробуйте пізніше."
        elif "400" in str(e):
            return "Не вдалося обробити запит. Перевірте чи вказан токен для роботи з ШІ моделлю."
        elif "connect" in str(e).lower():
            return "Не вдалося з'єднатися з сервером. Перевірте ваше інтернет-з'єднання та спробуйте пізніше."
        else:
            return "Сталася неочікувана помилка. Спробуйте ще раз."

class BaseView:
    def __init__(self, page: ft.Page):
        self.page = page

    def create_view(self, title):
        return ft.Container(
            content=ft.Column([
                ft.Text(title, style=ft.TextThemeStyle.HEADLINE_MEDIUM, font_family="Roboto"),
                ft.Text("Ця секція знаходиться в розробці", font_family="Roboto"),
            ]),
            alignment=ft.alignment.center,
            padding=20,
            expand=True,
        )
    

class VoiceRecognition:
    def __init__(self, on_result, on_error, on_system_message):
        self.config = load_config()
        self.recognizer = sr.Recognizer()
        self.is_listening = False
        self.language = self.config["language"]
        self.use_vosk = self.config["use_vosk"]
        self.preferred_recognition = self.config["preferred_recognition"]
        self.vosk_model = None
        self.audio_queue = queue.Queue()
        self.recognition_thread = None
        self.on_result = on_result
        self.on_error = on_error
        self.on_system_message = on_system_message
        self.vosk_models = self.config["vosk_models"]
        self.last_result = ""
        self.main_view = None
        self.processing_ai = False
        self.manually_selected_vosk = self.preferred_recognition == "vosk"
        self.last_internet_check = time.time()
        self.internet_check_interval = 5  # секунд
        self.stop_event = threading.Event()

    def change_language(self, language):
        self.language = language
        self.config["language"] = language
        save_config(self.config)
        if self.use_vosk:
            self.vosk_model = None
            model_path = self.vosk_models.get(self.language, "")
            if not model_path or not os.path.exists(model_path):
                self.on_error(f"Модель Vosk для {language} не знайдена. Перевірте шлях: {model_path}")
            else:
                try:
                    if self.main_view:
                        self.main_view.show_progress_ring("Завантаження моделі Vosk...")
                    self.vosk_model = Model(model_path)
                    if self.main_view:
                        self.main_view.hide_progress_ring()
                except Exception as e:
                    self.on_error(f"Помилка завантаження Vosk моделі: {str(e)}")
        self.restart_listening()

    def set_recognition_type(self, recognition_type):
        self.preferred_recognition = recognition_type
        self.manually_selected_vosk = (recognition_type == "vosk")
        self.config["preferred_recognition"] = recognition_type
        self.config["use_vosk"] = (recognition_type == "vosk")
        save_config(self.config)
        if recognition_type == "vosk":
            self.switch_to_vosk()
        else:
            self.switch_to_google()

    def switch_to_vosk(self, temporary=False):
        if not self.use_vosk:
            self.use_vosk = True
            if not temporary:
                self.preferred_recognition = "vosk"
                self.manually_selected_vosk = True
                self.config["preferred_recognition"] = "vosk"
            self.config["use_vosk"] = True
            save_config(self.config)
            self.on_system_message("Перемикання на офлайн-розпізнавання.")
            self.restart_listening()
        if self.main_view:
            self.main_view.update_mic_state()

    def switch_to_google(self):
        if self.use_vosk and self.preferred_recognition == "google" and not self.manually_selected_vosk:
            self.use_vosk = False
            self.config["use_vosk"] = False
            save_config(self.config)
            self.on_system_message("З'єднання з інтернетом відновлено. Перемикання на Google розпізнавання.")
            self.restart_listening()
        if self.main_view:
            self.main_view.update_mic_state()

    def listen_with_google(self):
        with sr.Microphone() as source:
            self.recognizer.adjust_for_ambient_noise(source, duration=0.2)
            while self.is_listening:
                try:
                    audio = self.recognizer.listen(source, timeout=0.7, phrase_time_limit=11)
                    self.audio_queue.put(audio)
                    threading.Thread(target=self.recognize_google, args=(audio,), daemon=True).start()
                except sr.WaitTimeoutError:
                    pass
                except Exception as e:
                    self.on_error(f"Помилка при прослуховуванні: {str(e)}")

    def recognize_google(self, audio):
        try:
            text = self.recognizer.recognize_google(audio, language=self.language)
            if text != self.last_result:
                self.last_result = text
                self.on_result(text)
        except sr.UnknownValueError:
            pass
        except sr.RequestError:
            self.switch_to_vosk()
        except Exception as e:
            self.on_error(f"Помилка при розпізнаванні: {str(e)}")

    def check_internet_connection(self):
        try:
            requests.get("http://www.google.com", timeout=3)
            return True
        except requests.ConnectionError:
            return False


    def monitor_internet_connection(self):
        while not self.stop_event.is_set():
            current_time = time.time()
            if current_time - self.last_internet_check >= self.internet_check_interval:
                has_internet = self.check_internet_connection()
                if not has_internet and self.preferred_recognition == "google":
                    self.switch_to_vosk(temporary=True)
                elif has_internet and self.use_vosk and self.preferred_recognition == "google" and not self.manually_selected_vosk:
                    self.switch_to_google()
                self.last_internet_check = current_time
            time.sleep(1)

    def start_listening(self):
        if not self.is_listening:
            self.is_listening = True
            self.last_result = ""
            self.audio_queue = queue.Queue()
            self.stop_event.clear()
            self.recognition_thread = threading.Thread(target=self.listen_in_background, daemon=True)
            self.recognition_thread.start()
            self.internet_monitor_thread = threading.Thread(target=self.monitor_internet_connection, daemon=True)
            self.internet_monitor_thread.start()
            if self.main_view:
                self.main_view.update_mic_state()

    def stop_listening(self):
        if self.is_listening:
            self.is_listening = False
            self.stop_event.set()
            if self.main_view:
                self.main_view.update_mic_state()

    def restart_listening(self):
        self.stop_listening()
        time.sleep(0.5)  # Даємо час на завершення потоків
        self.start_listening()

    def listen_in_background(self):
        try:
            if not self.use_vosk and self.check_internet_connection():
                self.listen_with_google()
            else:
                self.listen_with_vosk()
        finally:
            self.stop_listening()

    

    def listen_with_vosk(self):
        if self.vosk_model is None:
            model_path = self.vosk_models.get(self.language, "")
            if not model_path or not os.path.exists(model_path):
                self.on_error(f"Модель Vosk для {self.language} не знайдена. Перевірте шлях: {model_path}")
                return
            if self.main_view:
                self.main_view.show_progress_ring("Завантаження моделі Vosk...")
            self.vosk_model = Model(model_path)
            if self.main_view:
                self.main_view.hide_progress_ring()

        try:
            p = pyaudio.PyAudio()
            stream = p.open(format=pyaudio.paInt16, channels=1, rate=16000, input=True, frames_per_buffer=8000)
            stream.start_stream()

            rec = KaldiRecognizer(self.vosk_model, 16000)

            while self.is_listening:
                data = stream.read(4000, exception_on_overflow=False)
                if rec.AcceptWaveform(data):
                    result = json.loads(rec.Result())
                    text = result['text']
                    if text and text != self.last_result:
                        self.last_result = text
                        self.on_result(text)
                
                if hasattr(self, 'processing_ai') and self.processing_ai:
                    time.sleep(0.1)
                    continue

            stream.stop_stream()
            stream.close()
            p.terminate()
        except Exception as e:
            self.on_error(f"Помилка при використанні Vosk: {str(e)}")
        finally:
            self.is_listening = False

    def audio_callback(self, in_data, frame_count, time_info, status):
        if not self.processing_ai:
            self.audio_queue.put(in_data)
        return (None, pyaudio.paContinue)

    def set_vosk_model_path(self, language, path):
        if language in self.vosk_models:
            self.vosk_models[language] = path
            self.config["vosk_models"] = self.vosk_models
            save_config(self.config)
        else:
            self.on_error(f"Непідтримувана мова: {language}")

    
class WeatherView(BaseView):
    def build(self):
        return self.create_view("Погода")

class ChatView(BaseView):
    def build(self):
        return self.create_view("Чат")



class GesturesView(BaseView):
    def build(self):
        return self.create_view("Жести")

class CardsView(BaseView):
    def __init__(self, page: ft.Page):
        super().__init__(page)
        self.cards = []
        self.loaded_cards = load_saved_cards()
        self.open_extensions = []

        self.cards_view_download = ft.Column(
            spacing=10,
            height=500,  # Фіксована висота
            scroll=ft.ScrollMode.ALWAYS,
        )
        self.cards_view_loaded = ft.Column(
            spacing=10,
            height=500,  # Фіксована висота
            scroll=ft.ScrollMode.ALWAYS,
        )
        self.active_tab_index = 0

    def update_loaded_cards_view(self):
        self.cards_view_loaded.controls = [
            ft.ResponsiveRow(
                [ft.Column(col={"sm": 6, "md": 4, "lg": 3, "xl": 2}, controls=[self.create_flip_card(card, is_loaded=True)])
                 for card in self.loaded_cards]
            )
        ]
        save_loaded_cards(self.loaded_cards)
        self.page.update()

    def update_cards_view(self):
        self.cards_view_download.controls = [
            ft.ResponsiveRow(
                [ft.Column(col={"sm": 6, "md": 4, "lg": 3, "xl": 2}, controls=[card])
                 for card in self.cards]
            )
        ]
        self.page.update()

    async def load_functionality(self, card):
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(f"{SERVER_URL}/functionality/{card.id}") as response:
                    if response.status == 200:
                        functionality = await response.json()
                        card.content = functionality['content']
                        print(f"Завантажено функціональність для картки {card.id}")
                    else:
                        print(f"Помилка при завантаженні функціональності для картки {card.id}")
        except Exception as e:
            print(f"Помилка при завантаженні функціональності: {e}")
            raise

    def create_flip_card(self, card: Card, is_loaded=False):
        is_front = True

        def on_card_click(e):
            nonlocal is_front
            if is_loaded:
                self.open_new_tab(card)
            else:
                if is_front:
                    front_card.opacity = 0
                    back_card.opacity = 1
                else:
                    back_card.opacity = 0
                    front_card.opacity = 1
                is_front = not is_front
                self.page.update()

        def on_load_click(e):
            nonlocal is_loaded
            if not is_loaded:
                load_icon.icon = ft.icons.REFRESH
                load_icon.icon_color = ft.colors.BLUE_GREY_400
                self.page.update()

                async def load_and_update():
                    try:
                        await self.load_functionality(card)
                        self.loaded_cards.append(card)
                        self.cards.remove(card_container)
                        nonlocal is_loaded
                        is_loaded = True
                        update_icon_for_loaded_card()
                        save_loaded_cards(self.loaded_cards)
                        self.update_loaded_cards_view()
                        self.update_cards_view()
                    except Exception as ex:
                        print(f"Помилка при завантаженні: {ex}")
                    finally:
                        if not is_loaded:
                            load_icon.icon = ft.icons.DOWNLOAD
                            load_icon.icon_color = ft.colors.PRIMARY
                        self.page.update()

                asyncio.run(load_and_update())

        def on_delete_click(e):
            # Зберігаємо поточну вкладку
            current_tab = "Завантажені"  # Оскільки видалення відбувається з вкладки "Завантажені"
            
            self.loaded_cards.remove(card)
            save_loaded_cards(self.loaded_cards)
            self.cards.append(self.create_flip_card(card, is_loaded=False))
            
            # Видаляємо відповідну вкладку
            tab_to_remove = next((tab for tab in self.open_extensions if tab.text == card.functionality), None)
            if tab_to_remove:
                self.open_extensions.remove(tab_to_remove)
            
            self.update_loaded_cards_view()
            self.update_cards_view()
            self.update_tabs(active_tab=current_tab) 

        def update_icon_for_loaded_card():
            nonlocal load_icon
            load_icon.icon = ft.icons.DELETE
            load_icon.icon_color = ft.colors.RED_400
            load_icon.on_click = on_delete_click

        if is_loaded:
            load_icon = ft.IconButton(
                icon=ft.icons.DELETE,
                icon_color=ft.colors.RED_400,
                icon_size=30,
                on_click=on_delete_click,
            )
        else:
            load_icon = ft.IconButton(
                icon=ft.icons.DOWNLOAD,
                icon_color=ft.colors.PRIMARY,
                icon_size=30,
                on_click=on_load_click,
            )

        front_card = ft.Container(
            content=ft.Column(
                controls=[
                    ft.Stack(
                        controls=[
                            ft.Image(
                                src=card.image,
                                width=270,
                                height=340,
                                fit=ft.ImageFit.COVER,
                            ),
                            ft.Container(
                                content=ft.Row(
                                    controls=[
                                        ft.Text(
                                            card.functionality,
                                            size=18,
                                            color=ft.colors.WHITE,
                                            weight=ft.FontWeight.BOLD,
                                        )
                                    ],
                                    alignment=ft.MainAxisAlignment.CENTER,
                                ),
                                padding=ft.padding.symmetric(horizontal=10, vertical=5),
                                bgcolor=ft.colors.with_opacity(0.5, ft.colors.BLACK),
                                border_radius=ft.border_radius.all(10),
                                width=250,
                                height=30,
                                alignment=ft.alignment.top_center,
                            ),
                        ]
                    ),
                ],
                spacing=0,
            ),
            width=270,
            height=340,
            border_radius=ft.border_radius.all(10),
            shadow=ft.BoxShadow(blur_radius=10, color=ft.colors.BLACK),
            animate_opacity=300
        )

        back_card = ft.Container(
            content=ft.Column(
                controls=[
                    ft.Text("Опис картки:", size=24, color=ft.colors.WHITE),
                    ft.Text(
                        card.description,
                        size=18,
                        color=ft.colors.WHITE,
                        text_align=ft.TextAlign.CENTER
                    )
                ],
                alignment=ft.MainAxisAlignment.CENTER,
                horizontal_alignment=ft.CrossAxisAlignment.CENTER,
            ),
            width=270,
            height=370,
            bgcolor=ft.colors.BLUE,
            alignment=ft.alignment.center,
            border_radius=ft.border_radius.all(10),
            shadow=ft.BoxShadow(blur_radius=10, color=ft.colors.BLACK),
            opacity=0,
            animate_opacity=300
        )

        card_stack = ft.Stack(
            controls=[front_card, back_card]
        )

        load_icon_container = ft.Container(
            content=load_icon,
            alignment=ft.alignment.bottom_right,
            padding=10,
        )

        combined_stack = ft.Stack(
            controls=[
                card_stack,
                load_icon_container
            ]
        )

        card_container = ft.Container(
            content=combined_stack,
            width=250,  # Зменшуємо ширину картки
            height=350,  # Зменшуємо висоту картки
            on_click=on_card_click,
            border_radius=ft.border_radius.all(10),
            shadow=ft.BoxShadow(blur_radius=10, color=ft.colors.BLACK),
            alignment=ft.alignment.center
        )
        return card_container

    async def load_cards(self, _):
        try:
            # Завантаження карток із сервера
            cards_data = await fetch_cards()
            self.cards.clear()  # Очистити список карток

            # Додати картки до списку
            for card_data in cards_data:
                print(f"card_data: {card_data}")
                if 'date' in card_data:
                    del card_data['date']
                new_card = Card(**card_data)
                if new_card.id not in [card.id for card in self.loaded_cards]:
                    try:
                        self.cards.append(self.create_flip_card(new_card))
                    except TypeError as e:
                        print(f"Помилка при створенні картки: {e}")

            # Зачекати 2 секунди, щоб переконатися, що всі дані завантажені
            await asyncio.sleep(2) 

            # Оновити відображення завантажених карток
            self.update_cards_view() 
            self.update_tabs()  # Додати цей рядок
            self.page.update()

            # Завантаження завантажених карток
            self.loaded_cards = load_saved_cards()
            self.update_loaded_cards_view()
            self.page.update()
        except aiohttp.ClientConnectorError:
            self.show_error_message("Не вдалося підключитися до сервера. Перевірте підключення до інтернету або доступність сервера.")
        except Exception as e:
            self.show_error_message(f"Виникла помилка при завантаженні карток")#: {str(e)}
        finally:
            self.page.update()

    def show_error_message(self, message):
        self.cards_view_download.controls = [ft.Text(message, color=ft.colors.RED)]

    def load_cards_functionality(self):
        try:
            with open(CARDS_FUNCTIONALITY_FILE, 'r') as file:
                return json.load(file)
        except FileNotFoundError:
            return {}

    def open_new_tab(self, card):
        if not card.content:
            print(f"Error: Functionality for card {card.id} not found")
            return

        try:
            # Створюємо тимчасовий модуль
            spec = importlib.util.spec_from_loader(card.functionality, loader=None)
            module = importlib.util.module_from_spec(spec)
            
            # Виконуємо код
            exec(card.content, module.__dict__)
            
            # Отримуємо клас картки
            card_class = getattr(module, f"{card.functionality.capitalize()}Card")
            card_instance = card_class(self.page)
            
            content = card_instance.build()
            menu_items = card_instance.menu_items()

            new_tab = ft.Tab(
                text=card.functionality,
                content=ft.Column([
                    ft.Row([
                        ft.Text(card.functionality, style=ft.TextThemeStyle.HEADLINE_MEDIUM),
                        ft.PopupMenuButton(items=menu_items)
                    ]),
                    content
                ]),
            )
            
            self.open_extensions.append(new_tab)
            self.active_tab_index = len(self.open_extensions) + 1
            self.update_tabs()
        except Exception as e:
            print(f"Error opening new tab: {e}")

    def update_content(self):
        self.update_loaded_cards_view()
        self.update_cards_view()
        self.update_tabs()

    def update_tabs(self, active_tab=None):
        def create_tab_content(text, is_closable=False):
            content = [ft.Text(text, size=14, no_wrap=True, overflow=ft.TextOverflow.ELLIPSIS)]
            if is_closable:
                close_button = ft.IconButton(
                    icon=ft.icons.CLOSE,
                    icon_size=14,
                    on_click=lambda _: self.close_tab(text),
                    style=ft.ButtonStyle(
                        color={"": ft.colors.GREY_400, "hovered": ft.colors.RED_400},
                        shape=ft.CircleBorder(),
                    ),
                )
                content.append(close_button)
            return ft.Row(content, alignment=ft.MainAxisAlignment.CENTER, spacing=5, tight=True)

        tabs = [
            ft.Tab(
                tab_content=create_tab_content("Завантажити"),
                content=ft.Container(
                    content=ft.Column([self.cards_view_download]),
                    padding=10
                )
            ),
            ft.Tab(
                tab_content=create_tab_content("Завантажені"),
                content=ft.Container(
                    content=ft.Column([self.cards_view_loaded]),
                    padding=10
                )
            )
        ]
        if not self.open_extensions:
            self.active_tab_index = 0 
        
        for extension in self.open_extensions:
            tabs.append(ft.Tab(
                tab_content=create_tab_content(extension.text, True),
                content=extension.content
            ))
        
        # Встановлюємо активний індекс на основі переданої активної вкладки
        if active_tab:
            self.active_tab_index = next((i for i, tab in enumerate(tabs) if tab.tab_content.controls[0].value == active_tab), 0)
        else:
            # Якщо активна вкладка не передана, зберігаємо поточний індекс
            self.active_tab_index = min(self.active_tab_index, len(tabs) - 1)
        
        tab_bar = ft.Tabs(
            tabs=tabs,
            selected_index=self.active_tab_index,
            animation_duration=300,
            expand=1,
            tab_alignment=ft.TabAlignment.START,
        )
        
        self.content.content = ft.Column([
            ft.Container(
                content=tab_bar,
                padding=ft.padding.only(left=10, right=10, top=5),
                border=ft.border.only(bottom=ft.BorderSide(1, ft.colors.OUTLINE)),
            ),
            ft.Container(
                content=tabs[self.active_tab_index].content,
                expand=True,
                padding=10
            )
        ])
        self.page.update()

    def close_tab(self, tab_text):
        tab_to_remove = next((tab for tab in self.open_extensions if tab.text == tab_text), None)
        if tab_to_remove:
            index = self.open_extensions.index(tab_to_remove)
            self.open_extensions.remove(tab_to_remove)
            if self.active_tab_index > index + 1:  # +1 because of the two fixed tabs
                self.active_tab_index -= 1
            elif self.active_tab_index == index + 1:
                self.active_tab_index = min(self.active_tab_index, len(self.open_extensions) + 1)
            self.update_tabs()

    def build(self):
        self.content = ft.Container(
            content=ft.Tabs(
                tabs=[
                    ft.Tab(
                        text="Завантажити",
                        content=ft.Container(
                            content=self.cards_view_download,
                            padding=20,
                            expand=True
                        )
                    ),
                    ft.Tab(
                        text="Завантажені",
                        content=ft.Container(
                            content=self.cards_view_loaded,
                            padding=20,
                            expand=True
                        )
                    )
                ] + self.open_extensions,
                expand=1,
                selected_index=self.active_tab_index,
                animation_duration=300,
                divider_color=ft.colors.TRANSPARENT,
                indicator_color=ft.colors.PRIMARY,
                label_color=ft.colors.PRIMARY,
                unselected_label_color=ft.colors.GREY,
            ),
            expand=True
        )
        threading.Thread(target=lambda: asyncio.run(self.load_cards(None))).start()
        self.update_content()
        return self.content


class ExtrasView(BaseView):
    def build(self):
        return self.create_view("Додатково")
        
class MainView(ft.UserControl):
    def __init__(self, page: ft.Page):
        super().__init__()
        self.page = page
        self.messages = []
        self.voice_recognition = VoiceRecognition(
            on_result=self.on_voice_result,
            on_error=self.on_voice_error,
            on_system_message=self.on_system_message
        )
        self.voice_recognition.main_view = self
        self.progress_container = None
        self.dialog_box = None
        self.text_input = None
        self.send_button = None
        self.voice_input = None
        self.input_row = None
        self.command_processor = CommandProcessor(self)
        self.progress_lock = threading.Lock()
        self.loading_complete = threading.Event()
        self.animation_timer = None
        self.last_user_messages = []
        self.last_model_messages = []
        self.app_dir = os.path.dirname(os.path.abspath(__file__))

    def download_image(self, image_base64):
        try:
            image_data = base64.b64decode(image_base64)
            image = Image.open(BytesIO(image_data))
            image_path = os.path.join(self.app_dir, "generated_image.png")
            image.save(image_path)
            self.page.show_snack_bar(ft.SnackBar(content=ft.Text(f"Зображення збережено в: {image_path}")))
        except Exception as e:
            self.page.show_snack_bar(ft.SnackBar(content=ft.Text(f"Помилка завантаження зображення: {str(e)}")))

    def add_message(self, sender, content, content_type="text", color=None):
        self.messages.append({
            "sender": sender,
            "content": content,
            "type": content_type,
            "color": color
        })

        if sender == "попередні повідомленя від User":
            self.last_user_messages.append(content)
            if len(self.last_user_messages) > 2:
                self.last_user_messages.pop(0)
        elif sender == "ART":
            self.last_model_messages.append(content)
            if len(self.last_model_messages) > 2:
                self.last_model_messages.pop(0)

        self.update_message_colors()
        self.update_mic_state()
        self.page.update()

    def build(self):
        self.dialog_box = ft.ListView(
            expand=True,
            spacing=10,
            padding=20,
            auto_scroll=True,
        )

        self.text_input = ft.TextField(
            hint_text="Введіть команду...",
            border_radius=30,
            expand=True,
            on_submit=self.on_submit,
            focused_border_color=ft.colors.PRIMARY,
            focused_bgcolor=ft.colors.PRIMARY_CONTAINER,
        )

        self.send_button = ft.IconButton(
            icon=ft.icons.SEND,
            icon_color=ft.colors.PRIMARY,
            icon_size=30,
            on_click=self.on_send,
            tooltip="Надіслати",
        )

        self.voice_input = ft.IconButton(
            icon=ft.icons.MIC,
            icon_color=ft.colors.PRIMARY,
            icon_size=30,
            on_click=self.toggle_mic,
            tooltip="Голосове введення",
        )

        self.photo_button = ft.IconButton(
            icon=ft.icons.ADD_PHOTO_ALTERNATE,
            icon_color=ft.colors.PRIMARY,
            icon_size=30,
            on_click=self.upload_photo,
            tooltip="Додати фото(поки не працює)",
        )

        self.clear_button = ft.IconButton(
            icon=ft.icons.DELETE,
            icon_color=ft.colors.PRIMARY,
            icon_size=30,
            on_click=self.clear_chat,
            tooltip="Очистити чат",
        )

        self.input_row = ft.Container(
            content=ft.Row(
                [self.text_input, self.photo_button, self.send_button, self.voice_input, self.clear_button],
                alignment=ft.MainAxisAlignment.CENTER,
            ),
            padding=20,
            bgcolor=ft.colors.SURFACE,
            shadow=ft.BoxShadow(blur_radius=5, color=ft.colors.BLACK12, offset=ft.Offset(0, -3)),
        )

        main_content = ft.Column(
            [
                self.dialog_box,
                self.input_row,
            ],
            alignment=ft.MainAxisAlignment.END,
            horizontal_alignment=ft.CrossAxisAlignment.CENTER,
            expand=True,
        )

        self.update_message_colors()
        return main_content

    def on_submit(self, e):
        self.process_input(e.control.value)

    def on_send(self, _):
        self.process_input(self.text_input.value)

    def process_input(self, user_input):
        if user_input:
            self.add_message("User", user_input)
            self.command_processor.process_command(user_input)
            self.text_input.value = ""
            self.text_input.focus()
            self.update()

    def upload_photo(self, _):
        def on_result(e: ft.FilePickerResultEvent):
            if e.files:
                file_path = e.files[0].path
                self.add_message("User", file_path, content_type="image_local")
                self.page.update()

        file_picker = ft.FilePicker(on_result=on_result)
        self.page.overlay.append(file_picker)
        self.page.update()
        file_picker.pick_files(allow_multiple=False, allowed_extensions=["png", "jpg", "jpeg"])

    def clear_chat(self, e):
        self.messages.clear()
        self.update_message_colors()
        self.page.update()

    def update_message_colors(self):
        if self.dialog_box:
            self.dialog_box.controls.clear()
            new_controls = []
            for message in self.messages:
                if message["type"] == "text":
                    if self.page.theme_mode == ft.ThemeMode.DARK:
                        bgcolor = ft.colors.PRIMARY_CONTAINER if message["sender"] == "User" else ft.colors.SECONDARY_CONTAINER
                    else:
                        bgcolor = ft.colors.BACKGROUND if message["sender"] == "User" else ft.colors.SURFACE_VARIANT

                    border = ft.border.all(2, message["color"]) if message["color"] else None

                    new_controls.append(
                        ft.Container(
                            content=ft.Text(f"{message['sender']}: {message['content']}", selectable=True, font_family="Roboto"),
                            padding=15,
                            border_radius=15,
                            bgcolor=bgcolor,
                            border=border,
                            animate=ft.animation.Animation(300, ft.AnimationCurve.EASE_OUT),
                        )
                    )
                elif message["type"] == "image":
                    new_controls.append(
                        ft.CupertinoContextMenu(
                            content=ft.Image(
                                src_base64=message["content"],
                                width=245,
                                height=245,
                                fit=ft.ImageFit.COVER,
                                border_radius=15,
                            ),
                            actions=[
                                ft.CupertinoContextMenuAction(
                                    text="Завантажити",
                                    on_click=lambda e, img=message["content"]: self.download_image(img)
                                ),
                            ],
                        )
                    )
                elif message["type"] == "image_local":
                    new_controls.append(
                        ft.Container(
                            content=ft.Row([
                                ft.Image(
                                    src=message["content"],
                                    width=200,
                                    height=200,
                                    fit=ft.ImageFit.COVER,
                                    border_radius=15,
                                ),
                                ft.Container(expand=True)
                            ]),
                            padding=5,
                            border_radius=15,
                            bgcolor=ft.colors.TRANSPARENT,
                            animate=ft.animation.Animation(300, ft.AnimationCurve.EASE_OUT),
                        )
                    )
            self.dialog_box.controls.extend(new_controls)
            self.update()

    def update_mic_state(self):
        if self.voice_input:
            if self.voice_recognition.is_listening:
                self.voice_input.icon_color = ft.colors.ERROR
            else:
                self.voice_input.icon_color = ft.colors.PRIMARY
            self.page.update()

    def toggle_mic(self, e):
        if self.voice_recognition.is_listening:
            self.voice_recognition.stop_listening()
            self.add_message("Система", "Голосове введення вимкнено", color=ft.colors.BLUE)
        else:
            self.voice_recognition.start_listening()
            self.add_message("Система", "Голосове введення увімкнено", color=ft.colors.GREEN)

            def update_mic_state_after_start():
                self.update_mic_state()

            threading.Thread(target=update_mic_state_after_start, daemon=True).start()

        self.page.update()

    def on_voice_result(self, text):
        self.process_input(text)

    def on_voice_error(self, error):
        self.add_message("ART", f"Помилка: {error}")

    def on_system_message(self, message):
        self.add_message("система", message)

    def show_progress_ring(self, text):
        with self.progress_lock:
            if self.progress_container is None:
                progress_ring = ft.ProgressRing(
                    color=ft.colors.PRIMARY,
                    width=40,
                    height=40,
                    stroke_width=4,
                )

                progress_text = ft.Text(
                    "Модель офлайн розпізнавання завантажується",
                    color=ft.colors.PRIMARY,
                    size=16,
                    weight=ft.FontWeight.BOLD,
                )

                self.progress_container = ft.Container(
                    content=ft.Row(
                        [progress_ring, ft.Container(width=10), progress_text],
                        alignment=ft.MainAxisAlignment.CENTER,
                        vertical_alignment=ft.CrossAxisAlignment.CENTER,
                    ),
                    padding=10,
                    border_radius=10,
                    bgcolor=ft.colors.SURFACE_VARIANT,
                )

                self.page.add(self.progress_container)
                self.page.update()

                self.loading_complete.clear()

                def animate_progress():
                    while not self.loading_complete.is_set():
                        time.sleep(0.01)

                threading.Thread(target=animate_progress, daemon=True).start()

    def hide_progress_ring(self):
        self.loading_complete.set()
        with self.progress_lock:
            if self.progress_container:
                self.page.remove(self.progress_container)
                self.progress_container = None
                self.page.update()

    def on_model_loaded(self):
        self.hide_progress_ring()

class CommandType(Enum):
    WEBSITE = "Відкриття сайту"
    FILE = "Відкриття файлу"
    AUDIO = "Відтворення аудіо"
    SYSTEM = "Керування системою"
    DELAY = "Затримка"
    TEXT_TO_SPEECH = "Озвучення тексту" 

class SubCommand:
    def __init__(self, command_type: CommandType, data: Dict[str, Any]):
        self.command_type = command_type
        self.data = data

    def to_dict(self) -> Dict[str, Any]:
        return {
            "command_type": self.command_type.value,
            "data": self.data
        }

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'SubCommand':
        return cls(CommandType(data["command_type"]), data["data"])

class Command:
    def __init__(self, name: str, activation_word: str, icon: str):
        self.name = name
        self.activation_word = activation_word
        self.icon = icon
        self.sub_commands: List[SubCommand] = []
        self.is_favorite = False
        self.speaking_thread = None 

    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "activation_word": self.activation_word,
            "icon": self.icon,
            "sub_commands": [sc.to_dict() for sc in self.sub_commands],
            "is_favorite": self.is_favorite
        }

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'Command':
        command = cls(data["name"], data["activation_word"], data["icon"])
        command.sub_commands = [SubCommand.from_dict(sc) for sc in data["sub_commands"]]
        command.is_favorite = data.get("is_favorite", False)
        return command

    async def execute(self):
        for sub_command in self.sub_commands:
            await self.execute_sub_command(sub_command)

    async def execute_sub_command(self, sub_command: SubCommand):
        try:
            if sub_command.command_type == CommandType.WEBSITE:
                if "url" not in sub_command.data:
                    raise KeyError("Missing 'url' key in sub-command data")
                webbrowser.open(sub_command.data["url"])
            elif sub_command.command_type == CommandType.FILE:
                if "path" not in sub_command.data:
                    raise KeyError("Missing 'path' key in sub-command data")
                os.startfile(sub_command.data["path"])
            elif sub_command.command_type == CommandType.AUDIO:
                if "path" not in sub_command.data:
                    raise KeyError("Missing 'path' key in sub-command data")
                pygame.mixer.init()
                pygame.mixer.music.load(sub_command.data["path"])
                pygame.mixer.music.play()
            elif sub_command.command_type == CommandType.SYSTEM:
                if "action" not in sub_command.data:
                    raise KeyError("Missing 'action' key in sub-command data")
                action = sub_command.data["action"]
                if action == "shutdown":
                    os.system("shutdown /s /t 1")
                elif action == "restart":
                    os.system("shutdown /r /t 1")
                elif action == "sleep":
                    os.system("rundll32.exe powrprof.dll,SetSuspendState 0,1,0")
                elif action == "hibernate":
                    os.system("shutdown /h")
                elif action == "lock":
                    os.system("rundll32.exe user32.dll,LockWorkStation")
                elif action == "logout":
                    os.system("shutdown /l")
                elif action == "mute":
                    devices = AudioUtilities.GetSpeakers()
                    interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
                    volume = cast(interface, POINTER(IAudioEndpointVolume))
                    volume.SetMute(1, None)
                elif action == "unmute":
                    devices = AudioUtilities.GetSpeakers()
                    interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
                    volume = cast(interface, POINTER(IAudioEndpointVolume))
                    volume.SetMute(0, None)
                elif action == "volume_up":
                    devices = AudioUtilities.GetSpeakers()
                    interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
                    volume = cast(interface, POINTER(IAudioEndpointVolume))
                    current_volume = volume.GetMasterVolumeLevelScalar()
                    volume.SetMasterVolumeLevelScalar(min(1.0, current_volume + 0.1), None)
                elif action == "volume_down":
                    devices = AudioUtilities.GetSpeakers()
                    interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
                    volume = cast(interface, POINTER(IAudioEndpointVolume))
                    current_volume = volume.GetMasterVolumeLevelScalar()
                    volume.SetMasterVolumeLevelScalar(max(0.0, current_volume - 0.1), None)
                elif action == "empty_recycle_bin":
                    winshell.recycle_bin().empty(confirm=False, show_progress=False, sound=False)
            elif sub_command.command_type == CommandType.DELAY:
                if "seconds" not in sub_command.data:
                    raise KeyError("Missing 'seconds' key in sub-command data")
                await asyncio.sleep(float(sub_command.data["seconds"]))
            elif sub_command.command_type == CommandType.TEXT_TO_SPEECH:
                if "text" not in sub_command.data:
                    raise KeyError("Missing 'text' key in sub-command data")
                text = sub_command.data["text"]
                if text.strip():  # Перевірка, чи текст не є порожнім
                    self.speak(text)
                else:
                    print("Text for speech is empty or contains only whitespace.")
        except KeyError as e:
            print(f"Error executing sub-command: {e}")
        except Exception as e:
            print(f"Error executing sub-command: {e}")

    def speak(self, text):
        if len(text) < 500:
            if self.speaking_thread and self.speaking_thread.is_alive():
                sd.stop()
                self.speaking_thread.join()

            self.speaking_thread = threading.Thread(target=self._speak_thread, args=(text,))
            self.speaking_thread.start()

    def _speak_thread(self, text):
        config = load_config()
        if config.get("enable_tts", True):
            try:
                processed_text = preprocess_text(text)
                
                audio = model.apply_tts(text=processed_text,
                                        speaker=speaker,
                                        sample_rate=sample_rate,
                                        put_accent=True,
                                        put_yo=True)
                
                sd.play(audio.numpy(), sample_rate)
                sd.wait()
            except Exception as e:
                if str(e).strip():
                    self.main_view.add_message("ART", f"Помилка озвучування: {str(e)}")

class CommandsApp:
    def __init__(self, page: ft.Page):
        self.page = page
        self.commands: List[Command] = self.load_commands()
        self.current_view = None
        self.main_content = None

    def build(self):
        self.main_content = ft.Container(expand=True)
        nav_rail = ft.NavigationRail(
            selected_index=0,
            label_type=ft.NavigationRailLabelType.ALL,
            min_width=100,
            min_extended_width=300,
            group_alignment=-0.9,
            destinations=[
                ft.NavigationRailDestination(
                    icon=ft.icons.LIST,
                    selected_icon=ft.icons.LIST,
                    label="Список команд",
                ),
                ft.NavigationRailDestination(
                    icon=ft.icons.ADD,
                    selected_icon=ft.icons.ADD,
                    label="Створення команди",
                ),
                ft.NavigationRailDestination(
                    icon=ft.icons.HELP,
                    selected_icon=ft.icons.HELP,
                    label="Довідка",
                ),
            ],
            on_change=self.change_view,
        )

        return ft.Row(
            [
                nav_rail,
                ft.VerticalDivider(width=1),
                self.main_content,
            ],
            expand=True,
        )

    def change_view(self, e):
        if e.control.selected_index == 0:
            self.show_commands_list()
        elif e.control.selected_index == 1:
            self.show_command_creation()
        else:
            self.show_help()

    def show_commands_list(self):
        commands_list = CommandsList(self)
        self.current_view = commands_list
        self.main_content.content = commands_list
        self.page.update()

    def show_command_creation(self):
        command_creation = CommandsView(self)
        self.current_view = command_creation
        self.main_content.content = command_creation
        self.page.update()

    def show_help(self):
        help_dialog = ft.AlertDialog(
            modal=True,
            title=ft.Text("Довідка"),
            content=ft.Column(
                [
                    ft.Text("""
### 📚 Довідка для користувача програми керування командами

---

🚀 **Основні функції:**

1. **Створення команди:**
   - 📝 Введіть назву команди та слово для її активації.
   - 🎨 Оберіть іконку, щоб команда виглядала стильно!
   - 🧩 Додайте під-команди для виконання різних дій.

2. **Редагування команди:**
   - ✏️ Змінюйте назву, слово активації, іконку та під-команди у будь-який час.

3. **Виконання команди:**
   - ▶️ Натисніть "Play", щоб перевірити функціональність створеної команди. Це легко!

4. **Фільтрація команд:**
   - 🔍 Шукайте команди за назвою або словом активації, щоб швидко знайти потрібну.

---

🛠️ **Типи під-команд:**

1. **🌐 Відкриття сайту:**
   - Введіть URL сайту, і він миттєво відкриється!

2. **📂 Відкриття файлу:**
   - Вкажіть шлях до файлу, і він відкриється для вас.

3. **🎵 Відтворення аудіо:**
   - Введіть шлях до аудіофайлу, і він почне відтворюватися у фоновому режимі.
   - Аудіофайл, запущений через цю під-команду, не може бути зупинений, доки не завершиться відтворення.

4. **💻 Керування системою:**
   - Оберіть дію (наприклад, вимкнення чи перезавантаження), і система виконає її.

5. **⏱️ Затримка:**
   - Вкажіть кількість секунд, і програма почекає перед виконанням наступної дії.

6. **🗣️ Озвучення тексту:**
   - Введіть текст, і програма зачитає його вам.

---

💻 **Інтерфейс програми:**

- **📋 Список команд:** Тут ви бачите всі команди, можете їх шукати та виконувати.
- **➕ Створення команди:** Форма для створення або редагування команд.
- **ℹ️ Довідка:** Ця сторінка допоможе вам розібратися з усіма можливостями.

---

💡 **Поради:**

- 🧠 Використовуйте прості слова для активації, щоб їх було легко запам'ятати і щоб програма могла їх розпізнати.
- ⚠️ Використовуйте більше одного слова для активації, щоб уникнути випадкових активацій.
- 🧩 Плануйте під-команди заздалегідь, щоб створити послідовність дій, яка вам потрібна.

---

🛠️ **Помилки та проблеми:**

- ❗ Якщо щось пішло не так, перевірте дані у під-командах.
- 💬 У разі серйозних проблем, зверніться за допомогою до розробника.

---

😊 Дякуємо за використання нашої програми! Ми сподіваємося, що вона зробить ваші завдання легшими та цікавішими!
""")
                ],
                scroll=ft.ScrollMode.AUTO,
                expand=True,
            ),
            actions=[
                ft.TextButton("OK", on_click=lambda e: self.close_dialog()),
            ],
        )
        self.page.dialog = help_dialog
        self.page.dialog.open = True
        self.page.update()

    def close_dialog(self):
        self.page.dialog.open = False
        self.page.update()

    def load_commands(self):
        if os.path.exists(COMMANDS_FILE):
            with open(COMMANDS_FILE, "r", encoding="utf-8") as file:
                commands_data = json.load(file)
            return [Command.from_dict(cmd) for cmd in commands_data]
        return []

    def save_commands(self):
        with open(COMMANDS_FILE, "w", encoding="utf-8") as file:
            json.dump([cmd.to_dict() for cmd in self.commands], file, ensure_ascii=False, indent=4)

    def add_command(self, command):
        self.commands.append(command)
        self.save_commands()
        self.show_success_message(f"Команду '{command.name}' успішно додано")
        self.show_commands_list()  # Оновлення інтерфейсу

    def update_command(self, old_command, new_command):
        index = self.commands.index(old_command)
        self.commands[index] = new_command
        self.save_commands()
        self.show_success_message(f"Команду '{new_command.name}' успішно оновлено")
        self.show_commands_list()  # Оновлення інтерфейсу

    def delete_command(self, command):
        self.commands.remove(command)
        self.save_commands()
        self.show_success_message(f"Команду '{command.name}' успішно видалено")
        self.show_commands_list()  # Оновлення інтерфейсу

    def show_error_message(self, message):
        self.page.snack_bar = ft.SnackBar(
            content=ft.Text(message, color=ft.colors.WHITE),
            bgcolor=ft.colors.RED_400,
            action="OK",
            action_color=ft.colors.WHITE,
        )
        self.page.snack_bar.open = True
        self.page.update()

    def show_success_message(self, message):
        snack_bar = ft.SnackBar(
            content=ft.Text(message, color=ft.colors.WHITE),
            bgcolor=ft.colors.GREEN_400,
            action="OK",
            action_color=ft.colors.WHITE,
        )
        self.page.overlay.append(snack_bar)
        snack_bar.open = True
        self.page.update()

    def export_commands(self):
        with open("exported_commands.json", "w", encoding="utf-8") as file:
            json.dump([cmd.to_dict() for cmd in self.commands], file, ensure_ascii=False, indent=4)
        self.show_success_message("Команди успішно експортовано")

    def import_commands(self):
        if os.path.exists("exported_commands.json"):
            with open("exported_commands.json", "r", encoding="utf-8") as file:
                commands_data = json.load(file)
            self.commands.extend([Command.from_dict(cmd) for cmd in commands_data])
            self.save_commands()
            self.show_success_message("Команди успішно імпортовано")
        else:
            self.show_error_message("Файл для імпорту не знайдено")

class CommandsList(ft.Container):
    def __init__(self, app: CommandsApp):
        super().__init__()
        self.app = app
        self.commands_column = ft.Column(scroll=ft.ScrollMode.AUTO, expand=True)
        self.search_field = self.create_styled_textfield("Пошук команд", ft.icons.SEARCH)
        self.search_field.on_change = self.filter_commands
        self.content = self.build()

    def build(self):
        return ft.Container(
            content=ft.Column([
                ft.Text("Список команд", size=24, weight=ft.FontWeight.BOLD, color=ft.colors.BLUE),
                self.search_field,
                self.commands_column
            ]),
            padding=20,
            border_radius=10,
            bgcolor=ft.colors.WHITE,
            expand=True,
            border=ft.border.all(2, ft.colors.BLUE_400),
        )

    def did_mount(self):
        self.load_commands()

    def create_styled_textfield(self, label, icon):
        return ft.TextField(
            label=label,
            prefix_icon=icon,
            border_color=ft.colors.BLUE_400,
            border_radius=10,
            text_style=ft.TextStyle(color=ft.colors.BLUE_900),
            label_style=ft.TextStyle(color=ft.colors.BLUE_400),
            cursor_color=ft.colors.BLUE_400,
        )

    def build_command_view(self, command):
        return ft.Container(
            content=ft.Row([
                ft.Icon(name=command.icon, size=40, color=ft.colors.BLUE),
                ft.Column([
                    ft.Text(command.name, color=ft.colors.BLUE, size=16, weight=ft.FontWeight.BOLD),
                    ft.Text(f"Активація: {command.activation_word}", color=ft.colors.GREY, size=14),
                    ft.Text(f"Кількість під-команд: {len(command.sub_commands)}", color=ft.colors.GREY, size=14),
                ], expand=True),
                ft.Row([
                    ft.IconButton(ft.icons.PLAY_ARROW, on_click=self.execute_command_wrapper(command)),
                    ft.IconButton(ft.icons.EDIT, on_click=lambda _: self.edit_command(command)),
                    ft.IconButton(ft.icons.DELETE, on_click=lambda _: self.delete_command(command)),
                    ft.IconButton(ft.icons.STAR if command.is_favorite else ft.icons.STAR_BORDER, 
                                  on_click=lambda _: self.toggle_favorite(command))
                ]),
            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
            padding=10,
            border_radius=10,
            bgcolor=ft.colors.WHITE,
            border=ft.border.all(1, ft.colors.BLUE_200),
        )

    def execute_command_wrapper(self, command):
        def wrapper(_):
            try:
                loop = asyncio.get_event_loop()
            except RuntimeError:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
            
            if loop.is_running():
                loop.create_task(self.execute_command(command))
            else:
                loop.run_until_complete(self.execute_command(command))
        return wrapper

    async def execute_command(self, command):
        try:
            await command.execute()
            self.app.show_success_message(f"Виконується команда: {command.name}")
        except Exception as e:
            self.app.show_error_message(f"Помилка виконання команди {command.name}: {str(e)}")

    def edit_command(self, command):
        self.app.show_command_creation()
        self.app.current_view.load_command(command)

    def delete_command(self, command):
        def confirm_delete(e):
            if e.control.data == "yes":
                self.app.delete_command(command)
                self.load_commands()
            self.app.page.dialog.open = False
            self.app.page.update()

        self.app.page.dialog = ft.AlertDialog(
            modal=True,
            title=ft.Text("Підтвердження видалення"),
            content=ft.Text(f"Ви впевнені, що хочете видалити команду '{command.name}'?"),
            actions=[
                ft.TextButton("Так", on_click=confirm_delete, data="yes"),
                ft.TextButton("Ні", on_click=confirm_delete, data="no"),
            ],
        )
        self.app.page.dialog.open = True
        self.app.page.update()

    def load_commands(self):
        self.commands_column.controls.clear()
        sorted_commands = sorted(self.app.commands, key=lambda x: (not x.is_favorite, x.name.lower()))
        for command in sorted_commands:
            self.commands_column.controls.append(self.build_command_view(command))
        if self.page:
            self.page.update()

    def toggle_favorite(self, command):
        command.is_favorite = not command.is_favorite
        self.app.save_commands()
        self.load_commands() 
        
    def filter_commands(self, e):
        search_query = e.control.value.lower()
        filtered_commands = [cmd for cmd in self.app.commands if search_query in cmd.name.lower() or search_query in cmd.activation_word.lower()]
        self.commands_column.controls.clear()
        for command in filtered_commands:
            self.commands_column.controls.append(self.build_command_view(command))
        if self.page:
            self.page.update()

class CommandsView(ft.Container):
    def __init__(self, app: CommandsApp):
        super().__init__()
        self.app = app
        self.editing_command = None
        self.selected_icon = None
        self.init_ui_elements()
        self.content = self.build()

    def init_ui_elements(self):
        self.command_name = self.create_styled_textfield("Назва команди", ft.icons.CREATE)
        self.activation_word = self.create_styled_textfield("Активаційне слово/слова", ft.icons.VOICE_CHAT)
        self.sub_commands_column = ft.Column(spacing=10)
        self.icon_grid = self.create_icon_grid()
        self.save_button = ft.ElevatedButton(
            "Зберегти команду",
            on_click=self.save_command,
            style=ft.ButtonStyle(
                color=ft.colors.WHITE,
                bgcolor=ft.colors.BLUE_400,
                shape=ft.RoundedRectangleBorder(radius=10),
            ),
        )

    def build(self):
        content = ft.Column([
            ft.Text("Створення команди", size=24, weight=ft.FontWeight.BOLD, color=ft.colors.BLUE),
            self.command_name,
            self.activation_word,
            ft.Text("Під-команди:", size=16, color=ft.colors.BLUE),
            self.sub_commands_column,
            ft.ElevatedButton(
                "Додати під-команду",
                on_click=self.add_sub_command,
                style=ft.ButtonStyle(
                    color=ft.colors.WHITE,
                    bgcolor=ft.colors.BLUE_400,
                    shape=ft.RoundedRectangleBorder(radius=10),
                ),
            ),
            ft.Text("Оберіть іконку:", size=16, color=ft.colors.BLUE),
            self.icon_grid,
        ], spacing=20, scroll=ft.ScrollMode.AUTO)

        return ft.Container(
            content=ft.Column([
                ft.Container(content=content, expand=True),
                self.save_button,
            ]),
            padding=20,
            border_radius=10,
            bgcolor=ft.colors.WHITE,
            expand=True,
            border=ft.border.all(2, ft.colors.BLUE_400),
        )

    def create_styled_textfield(self, label, icon):
        return ft.TextField(
            label=label,
            prefix_icon=icon,
            border_color=ft.colors.BLUE_400,
            border_radius=10,
            text_style=ft.TextStyle(color=ft.colors.BLUE_900),
            label_style=ft.TextStyle(color=ft.colors.BLUE_400),
            cursor_color=ft.colors.BLUE_400,
        )

    def create_icon_grid(self):
        icon_grid = ft.GridView(
            expand=1,
            runs_count=5,
            max_extent=50,
            child_aspect_ratio=1,
            spacing=5,
            run_spacing=5,
        )

        icons = [
            ft.icons.HOME, ft.icons.SETTINGS, ft.icons.ADD_TASK, ft.icons.BOOKMARK, ft.icons.FAVORITE, ft.icons.SEARCH,
            ft.icons.EMAIL, ft.icons.CALL, ft.icons.MESSAGE, ft.icons.NOTIFICATIONS, ft.icons.CALENDAR_TODAY,
            ft.icons.SHOPPING_CART, ft.icons.PAYMENT, ft.icons.ACCOUNT_BALANCE, ft.icons.DIRECTIONS_BUS, ft.icons.FLIGHT,
            ft.icons.HOTEL, ft.icons.RESTAURANT, ft.icons.LOCAL_GROCERY_STORE, ft.icons.FITNESS_CENTER, ft.icons.HEALTH_AND_SAFETY,
            ft.icons.BOOK_ONLINE, ft.icons.EVENT, ft.icons.MUSIC_NOTE, ft.icons.MOVIE, ft.icons.PHOTO_CAMERA,
            ft.icons.VIDEOGAME_ASSET, ft.icons.COLOR_LENS, ft.icons.LANGUAGE, ft.icons.PUBLIC,
            ft.icons.LOCATION_ON, ft.icons.NAVIGATE_NEXT, ft.icons.NAVIGATE_BEFORE, ft.icons.ZOOM_IN,
            ft.icons.ZOOM_OUT, ft.icons.FULLSCREEN, ft.icons.FULLSCREEN_EXIT, ft.icons.LOCK, ft.icons.LOCK_OPEN,
            ft.icons.PASSWORD, ft.icons.SECURITY, ft.icons.VPN_KEY, ft.icons.PHONE_IPHONE, ft.icons.PHONE_ANDROID,
            ft.icons.COMPUTER, ft.icons.LAPTOP, ft.icons.DESKTOP_WINDOWS, ft.icons.DESKTOP_MAC, ft.icons.MONITOR,
            ft.icons.SWITCH_ACCOUNT, ft.icons.SYNC, ft.icons.SYNC_DISABLED, ft.icons.STORAGE, ft.icons.MEMORY, ft.icons.BATTERY_FULL,
            ft.icons.BATTERY_ALERT, ft.icons.POWER, ft.icons.SETTINGS_POWER, ft.icons.TIMER, ft.icons.TIMER_OFF,
        ]


        for icon in icons:
            icon_button = ft.IconButton(
                icon=icon,
                on_click=lambda _, i=icon: self.select_icon(i),
                style=ft.ButtonStyle(
                    color=ft.colors.BLUE,
                    bgcolor={ft.MaterialState.HOVERED: ft.colors.BLUE_100},
                )
            )
            icon_grid.controls.append(icon_button)
        return icon_grid

    def select_icon(self, icon):
        self.selected_icon = icon
        for icon_button in self.icon_grid.controls:
            icon_button.style.bgcolor = ft.colors.BLUE_100 if icon_button.icon == icon else ft.colors.TRANSPARENT
        if self.page:
            self.page.update()

    def add_sub_command(self, e):
        sub_command_view = SubCommandView(self, on_remove=self.remove_sub_command)
        self.sub_commands_column.controls.append(sub_command_view)
        if self.page:
            self.page.update()

    def remove_sub_command(self, sub_command_view):
        self.sub_commands_column.controls.remove(sub_command_view)
        if self.page:
            self.page.update()

    def save_command(self, e):
        name = self.command_name.value
        activation_word = self.activation_word.value

        if not name or not activation_word:
            self.app.show_error_message("Будь ласка, заповніть назву команди та активаційне слово")
            return

        if not self.selected_icon:
            self.app.show_error_message("Будь ласка, оберіть іконку")
            return

        if not self.sub_commands_column.controls:
            self.app.show_error_message("Додайте хоча б одну під-команду")
            return

        new_command = Command(name, activation_word, self.selected_icon)
        new_command.sub_commands = self.get_sub_commands()

        if self.editing_command:
            self.app.update_command(self.editing_command, new_command)
            self.save_button.text = "Оновити команду"
        else:
            self.app.add_command(new_command)
            self.save_button.text = "Зберегти команду"

        self.clear_form()
        self.app.show_commands_list()

    def get_sub_commands(self):
        return [sc.get_sub_command() for sc in self.sub_commands_column.controls]

    def clear_form(self):
        self.command_name.value = ""
        self.activation_word.value = ""
        self.sub_commands_column.controls.clear()
        self.selected_icon = None
        for icon_button in self.icon_grid.controls:
            icon_button.style.bgcolor = ft.colors.TRANSPARENT
        self.editing_command = None
        if self.page:
            self.page.update()

    def load_command(self, command):
        self.editing_command = command
        self.command_name.value = command.name
        self.activation_word.value = command.activation_word
        self.sub_commands_column.controls.clear()
        for sub_command in command.sub_commands:
            sub_command_view = SubCommandView(self, on_remove=self.remove_sub_command)
            self.sub_commands_column.controls.append(sub_command_view)
            sub_command_view.load_sub_command(sub_command)
        self.select_icon(command.icon)
        if self.page:
            self.page.update()

class SubCommandView(ft.Container):
    def __init__(self, parent, on_remove):
        super().__init__()
        self.parent = parent
        self.on_remove = on_remove
        self.command_type = None
        self.data_fields = ft.Column()
        self.content = self.build()

    def build(self):
        return ft.Container(
            content=ft.Column([
                ft.Row([
                    ft.Dropdown(
                        label="Тип команди",
                        options=[ft.dropdown.Option(ct.value) for ct in CommandType],
                        width=250,
                        on_change=self.on_command_type_change
                    ),
                    ft.IconButton(ft.icons.REMOVE, on_click=self.remove_self)
                ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
                self.data_fields
            ]),
            padding=10,
            border_radius=10,
            bgcolor=ft.colors.BLUE_50,
        )

    def on_command_type_change(self, e):
        self.command_type = CommandType(e.control.value)
        self.data_fields.controls.clear()
        self.data_fields.controls.append(self.get_additional_data_fields(self.command_type))
        if self.page:
            self.page.update()

    def get_additional_data_fields(self, command_type: CommandType):
        if command_type == CommandType.WEBSITE:
            return self.create_styled_textfield("URL", ft.icons.LINK)
        elif command_type in [CommandType.FILE, CommandType.AUDIO]:
            return self.create_styled_textfield("Шлях", ft.icons.FOLDER)
        elif command_type == CommandType.SYSTEM:
            return ft.Dropdown(
                label="Системна дія",
                options=[
                    ft.dropdown.Option("shutdown", "Вимкнення"),
                    ft.dropdown.Option("restart", "Перезавантаження"),
                    ft.dropdown.Option("sleep", "Сон"),
                    ft.dropdown.Option("hibernate", "Гібернація"),
                    ft.dropdown.Option("lock", "Блокування"),
                    ft.dropdown.Option("logout", "Вихід з системи"),
                    ft.dropdown.Option("mute", "Вимкнення звуку"),
                    ft.dropdown.Option("unmute", "Увімкнення звуку"),
                    # ft.dropdown.Option("volume_up", "Збільшення гучності"),
                    # ft.dropdown.Option("volume_down", "Зменшення гучності"),
                    ft.dropdown.Option("empty_recycle_bin", "Очистити кошик"),
                ],
                width=250
            )
        elif command_type == CommandType.DELAY:
            return self.create_styled_textfield("Затримка (секунди)", ft.icons.TIMER)
        elif command_type == CommandType.TEXT_TO_SPEECH:
            return self.create_styled_textfield("Текст для озвучення", ft.icons.RECORD_VOICE_OVER)
        return ft.Container()

    def create_styled_textfield(self, label, icon):
        return ft.TextField(
            label=label,
            prefix_icon=icon,
            border_color=ft.colors.BLUE_400,
            border_radius=10,
            text_style=ft.TextStyle(color=ft.colors.BLUE_900),
            label_style=ft.TextStyle(color=ft.colors.BLUE_400),
            cursor_color=ft.colors.BLUE_400,
        )

    def remove_self(self, e):
        self.on_remove(self)

    def get_sub_command(self):
        data = {}
        if self.command_type == CommandType.WEBSITE:
            data["url"] = self.data_fields.controls[0].value
        elif self.command_type in [CommandType.FILE, CommandType.AUDIO]:
            data["path"] = self.data_fields.controls[0].value
        elif self.command_type == CommandType.SYSTEM:
            data["action"] = self.data_fields.controls[0].value
        elif self.command_type == CommandType.DELAY:
            data["seconds"] = self.data_fields.controls[0].value
        elif self.command_type == CommandType.TEXT_TO_SPEECH:
            data["text"] = self.data_fields.controls[0].value
        return SubCommand(self.command_type, data)

    def load_sub_command(self, sub_command):
        self.command_type = sub_command.command_type
        self.data_fields.controls.clear()
        additional_fields = self.get_additional_data_fields(self.command_type)
        self.data_fields.controls.append(additional_fields)
        
        if self.command_type == CommandType.WEBSITE:
            additional_fields.value = sub_command.data.get("url", "")
        elif self.command_type in [CommandType.FILE, CommandType.AUDIO]:
            additional_fields.value = sub_command.data.get("path", "")
        elif self.command_type == CommandType.SYSTEM:
            additional_fields.value = sub_command.data.get("action", "")
        elif self.command_type == CommandType.TEXT_TO_SPEECH:
            additional_fields.value = sub_command.data.get("text", "")
        elif self.command_type == CommandType.DELAY:
            additional_fields.value = sub_command.data.get("seconds", "")
        self.parent.update()

class CommandProcessor:
    def __init__(self, main_view):
        self.main_view = main_view
        self.config = load_config()
        self.speaking_thread = None
        self.AI_URL = f"https://api-inference.huggingface.co/models/{self.config.get('ai_model', 'mistralai/Mixtral-8x7B-Instruct-v0.1')}"
        self.FLUX_URL = f"https://api-inference.huggingface.co/models/{self.config.get('flux_model', 'black-forest-labs/FLUX.1-dev')}"
        self.headers = {"Authorization": f"Bearer {self.config.get('hf_api_key', '')}"}
        self.last_user_messages = []
        self.last_model_messages = []
        self.groq_client = setup_groq_client()
        self.history = []  # Історія повідомлень
        self.translator = GoogleTranslator()  # Ініціалізуємо GoogleTranslator
        self.commands = self.load_commands()  # Завантаження команд

    def load_commands(self):
        # Завантаження команд з файлу або іншого джерела
        if os.path.exists(COMMANDS_FILE):
            with open(COMMANDS_FILE, "r", encoding="utf-8") as file:
                commands_data = json.load(file)
            return [Command.from_dict(cmd) for cmd in commands_data]
        return []

    def update_commands(self):
        self.commands = self.load_commands()

    def update_headers(self):
        self.headers = {"Authorization": f"Bearer {self.config.get('hf_api_key', '')}"}

    def add_message(self, role, content):
        try:
            # translated_content = self.translate(content, 'uk', 'en')
            message = f'{role}: {content}'
            if message not in self.history:
                self.history.append(message)
                if len(self.history) > 6:
                    self.history = self.history[-6:]
                print(self.history)
        except Exception as e:
            logging.error(f"Помилка при додаванні повідомлення: {str(e)}")

    def process_command(self, command):
        self.update_commands() # Додане оновлення команд перед обробкою
        command = command.lower()
        print(f"Отримана команда: {command}")
        print(f"Використання ШІ: {self.config.get('use_ai', False)}")
        command_executed = False

        for cmd in self.commands:
            if cmd.activation_word.lower() in command:
                print(f"Знайдено відповідну команду: {cmd.name}")
                try:
                    # Перевіряємо, чи команда містить підкоманду озвучування тексту
                    has_text_to_speech = any(sub_cmd.command_type == CommandType.TEXT_TO_SPEECH for sub_cmd in cmd.sub_commands)
                    
                    # Створюємо новий цикл подій, якщо його немає
                    try:
                        loop = asyncio.get_running_loop()
                    except RuntimeError:
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)

                    with ThreadPoolExecutor() as executor:
                        loop.run_in_executor(executor, lambda: asyncio.run(cmd.execute()))
                    
                    # Озвучуємо повідомлення тільки якщо команда не містить підкоманду озвучування тексту
                    if not has_text_to_speech:
                        response = f"Виконується команда: {cmd.name}"
                        self.speaknwrite(response)
                    else:
                        # Просто додаємо повідомлення в інтерфейс без озвучування
                        self.main_view.add_message("ART", f"Виконується команда: {cmd.name}")
                except Exception as e:
                    logging.error(f"Помилка виконання команди {cmd.name}: {str(e)}")
                    response = f"Виникла помилка при виконанні команди {cmd.name}"
                    self.speaknwrite(response)
                command_executed = True
                break

        if not command_executed:  # Перевіряємо прапорець перед обробкою інших команд
           
            if "напиши текст" in command:
                text_to_write = command.split("напиши текст", 1)[1].strip()
                if text_to_write:
                    keyboard.write(text_to_write)
                    response = "Текст введено."
                    self.speak(response)
                else:
                    response = "Будь ласка, вкажіть текст після команди 'напиши текст'."
                    self.speak(response)

            elif "згенеруй зображення" in command:
                self.generate_image(command)
                return
            elif self.config.get("use_ai", False):
                    if not self.check_internet_connection():
                        self.write("Вибачте, немає підключення до інтернету. Не вдалося отримати відповідь від ШІ.")
                        return

                    self.last_user_messages.append(f"User: {command}")
                    if len(self.last_user_messages) > 3:
                        self.last_user_messages.pop(0)

                    context = "\n".join(self.last_user_messages + self.last_model_messages)

                    if self.config.get("text_generation_model") == "groq":
                        response = self.process_groq_command(command)
                    else:
                        response = self.process_custom_command(command, context)

                    print(f"AI відповідь: {response}")

                    self.last_model_messages.append(f"ART: {response}")
                    if len(self.last_model_messages) > 3:
                        self.last_model_messages.pop(0)

                    # Видаляємо виклик process_response звідси, оскільки він вже викликається в handle_request
                    self.speaknwrite(response)

                    # Збереження історії повідомлень
                    self.add_message("user", command)
                    self.add_message("assistant", response)
            else:
                response = "Вибачте, я не розумію цю команду."

            print(f"Відповідь: {response}")

    def process_groq_command(self, command):
        messages = [
            {"role": "system", "content": generate_system_prompt()},
            {"role": "user", "content": command}
        ]
        response = handle_request(messages, self.history, self)  # Додайте self як аргумент

        # translated_text = translate_text(response, 'en', 'uk') 
        
        # Оновлюємо історію
        self.add_message("user", command)
        self.add_message("assistant", response)
        
        return response.replace("Помічник:".lower(), "", 1).strip()

    def process_custom_command(self, command, context):
        parameters = {"max_new_tokens": 512 if "есе" in command else 100}
        return self.get_ai_response(command, context, parameters)  # Ensure context is passed here

    def reload_ai_config(self):
        self.config = load_config()
        self.AI_URL = f"https://api-inference.huggingface.co/models/{self.config.get('ai_model', 'mistralai/Mixtral-8x7B-Instruct-v0.1')}"
        self.FLUX_URL = f"https://api-inference.huggingface.co/models/{self.config.get('flux_model', 'black-forest-labs/FLUX.1-dev')}"
        self.update_headers()  # Оновлюємо заголовки
        self.groq_client = setup_groq_client()
    
    def generate_image(self, command):
        prompt = command.replace("згенеруй зображення", "").strip()
        if not prompt:
            self.write("Будь ласка, вкажіть опис для генерації зображення.")
            return

        translated_prompt = self.translate(prompt, "uk", "en")
        print(f"Використовується модель FLUX: {self.config.get('flux_model', 'black-forest-labs/FLUX.1-dev')}")
        flux_output = self.flux_query(translated_prompt)

        if isinstance(flux_output, bytes):
            image_base64 = base64.b64encode(flux_output).decode()
            self.main_view.add_message("ART", image_base64, content_type="image")
        else:
            self.write(f"Помилка при отриманні відповіді від FLUX: {flux_output}")

    def query(self, url, payload, max_retries=5, retry_delay=20):
        for attempt in range(max_retries):
            response = requests.post(url, headers=self.headers, json=payload)

            if response.status_code == 200:
                if 'image/' in response.headers.get('Content-Type', ''):
                    return response.content

                try:
                    return response.json()
                except ValueError:
                    print("Отримано бінарні дані, але вони не є зображенням.")
                    return {"error": "Невідомий формат відповіді."}

            if response.status_code == 503 and "is currently loading" in response.text:
                print(f"Модель завантажується. Спроба {attempt + 1} з {max_retries}. Очікування {retry_delay} секунд...")
                time.sleep(retry_delay)
            else:
                print(f"Отримано невідому помилку: {response.status_code}")
                return {"error": f"Запит завершився помилкою: {response.status_code}"}

        return {"error": "Модель не завантажилась після кількох спроб."}

    def translate(self, text, source_lang, target_lang):
        try:
            translation = self.translator.translate(text, source=source_lang, target=target_lang)
            return translation
        except Exception as e:
            logging.error(f"Помилка при перекладі з {source_lang} на {target_lang}: {str(e)}")
            return text

    def flux_query(self, prompt):
        payload = {
            "inputs": prompt,
        }
        return self.query(self.FLUX_URL, payload)

    def check_internet_connection(self):
        try:
            response = requests.get("http://www.google.com", timeout=5)
            response.raise_for_status()
            return True
        except requests.RequestException:
            return False

    def get_ai_response(self, text, context, parameters):
        print(f"Використовується модель ШІ: {self.config.get('ai_model', 'mistralai/Mixtral-8x7B-Instruct-v0.1')}")
        ai_output = self.ai_query(text, parameters, context)  # Query with the original Ukrainian text
        
        if isinstance(ai_output, list) and len(ai_output) > 0 and "generated_text" in ai_output[0]:
            en_response = ai_output[0]["generated_text"]
            
            # Видалення фрагменту між <s>[INST] та [/INST]
            cleaned_response = re.sub(r'<s>\[INST\].*?\[/INST\]', '', en_response, flags=re.DOTALL)
            
            # Переклад відповіді на українську мову
            uk_response = translate_text(cleaned_response, "en", "uk")
            print(f"AI відповідь: {uk_response}")
            return uk_response.strip() 
        else:
            return None     

    def ai_query(self, text, parameters, context):
        system_prompt = "Системний промпт розробника: ви корисний асистент. Розмовляйте з використанням емодзі. Ваше ім'я ART"
        payload = {
            "inputs": f"<s>[INST] {system_prompt}\n\n{context}\n\n{text} [/INST]",  # Додаємо контекст
            "parameters": parameters
        }
        try:
            return self.query(self.AI_URL, payload)
        except Exception as e:
            if "NETWORK_ERROR" in str(e):
                self.write("Вибачте, немає підключення до інтернету. Не вдалося отримати відповідь від ШІ.")
            else:
                self.write(f"Помилка при запиті до моделі ШІ: {str(e)}")
            return {"error": "Не вдалося зв'язатися з моделлю ШІ."}
        
    def speak(self, text):
        if len(text) < 500:
            if self.speaking_thread and self.speaking_thread.is_alive():
                # Якщо попередній потік ще активний, зупиняємо його
                sd.stop()
                self.speaking_thread.join()

            self.speaking_thread = threading.Thread(target=self._speak_thread, args=(text,))
            self.speaking_thread.start()

    def _speak_thread(self, text):
        config = load_config()
        if config.get("enable_tts", True):
            try:
                processed_text = preprocess_text(text)
                
                audio = model.apply_tts(text=processed_text,
                                        speaker=speaker,
                                        sample_rate=sample_rate,
                                        put_accent=True,
                                        put_yo=True)
                
                sd.play(audio.numpy(), sample_rate)
                sd.wait()
            except Exception as e:
                if str(e).strip():
                    self.main_view.add_message("ART", f"Помилка озвучування: {str(e)}")
       
    def speaknwrite(self, text):
        self.main_view.add_message("ART", text)
        config = load_config()
        if config.get("enable_tts", True):
            self.speak(text)

    def write(self, text):
        self.main_view.add_message("ART", text)

class SettingsView(BaseView):
    def __init__(self, page: ft.Page, voice_recognition: VoiceRecognition, command_processor: CommandProcessor):
        super().__init__(page)
        self.voice_recognition = voice_recognition
        self.config = load_config()
        self.current_settings_tab = None
        self.command_processor = command_processor

        self.language_dropdown = ft.Dropdown(
            label="Мова розпізнавання",
            options=[
                ft.dropdown.Option("uk-UA", "Українська"),
                ft.dropdown.Option("en-US", "English"),
                ft.dropdown.Option("ru-RU", "Русский"),
            ],
            value=self.voice_recognition.language,
            on_change=self.change_language,
            width=300,
            border_color=ft.colors.PRIMARY,
            focused_border_color=ft.colors.SECONDARY,
        )

        self.recognition_type = ft.RadioGroup(
            content=ft.Column([
                ft.Radio(value="google", label="Google Speech Recognition"),
                ft.Radio(value="vosk", label="Vosk (офлайн)"),
            ]),
            value="google" if not self.voice_recognition.use_vosk else "vosk",
            on_change=self.change_recognition_type,
        )

        self.vosk_model_paths = {
            "uk-UA": ft.TextField(label="Шлях до української Vosk моделі", value=self.voice_recognition.vosk_models["uk-UA"], width=300, border_color=ft.colors.PRIMARY),
            "en-US": ft.TextField(label="Шлях до англійської Vosk моделі", value=self.voice_recognition.vosk_models["en-US"], width=300, border_color=ft.colors.PRIMARY),
            "ru-RU": ft.TextField(label="Шлях до російської Vosk моделі", value=self.voice_recognition.vosk_models["ru-RU"], width=300, border_color=ft.colors.PRIMARY),
        }

        self.enable_tts = ft.Checkbox(
            label="Увімкнути озвучування команд",
            value=load_config().get("enable_tts", True),
            on_change=self.toggle_tts
        )

        self.use_ai = ft.Checkbox(
            label="Використовувати ШІ для відповідей",
            value=self.config.get("use_ai", False),
            on_change=self.toggle_ai
        )

        self.ai_model = ft.Dropdown(
            label="Назва моделі ШІ",
            options=[
                ft.dropdown.Option("mistralai/Mixtral-8x7B-Instruct-v0.1", "Mistral 7B"),
                ft.dropdown.Option("fireworks/llama-v2-7b-chat", "Fireworks Llama 7B"),
                ft.dropdown.Option("fireworks/llama-v2-13b-chat", "Fireworks Llama 13B"),
                ft.dropdown.Option("fireworks/llama-v2-30b-chat", "Fireworks Llama 30B"),
                ft.dropdown.Option("fireworks/llama-v2-65b-chat", "Fireworks Llama 65B"),
                ft.dropdown.Option("fireworks/llama-v2-70b-chat", "Fireworks Llama 70B"),
            ],
            value=self.config.get("ai_model", "mistralai/Mixtral-8x7B-Instruct-v0.1"),
            width=300,
            border_color=ft.colors.PRIMARY,
        )

        self.hf_api_key = ft.TextField(
            label="Ключ Hugging Face API",
            value=self.config.get("hf_api_key", ""),
            width=300,
            border_color=ft.colors.PRIMARY,
            password=True  # Щоб приховати ключ
        )

        self.groq_api_token = ft.TextField(
            label="Системний токен",
            value=self.config.get("groq_api_token", "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"),
            width=300,
            border_color=ft.colors.PRIMARY,
            password=True  # To hide the token
        )

        self.flux_model = ft.TextField(
            label="Назва моделі FLUX",
            value=self.config.get("flux_model", "black-forest-labs/FLUX.1-dev"),
            width=300,
            border_color=ft.colors.PRIMARY,
        )

        save_button = ft.ElevatedButton(
            "Зберегти налаштування", 
            on_click=self.save_settings,
            style=ft.ButtonStyle(
                color={ft.MaterialState.DEFAULT: ft.colors.WHITE},
                bgcolor={ft.MaterialState.DEFAULT: ft.colors.PRIMARY},
                padding=10,
            )
        )

        self.voice_recognition_settings = ft.Column([
            ft.Text("Налаштування розпізнавання голосу", style=ft.TextThemeStyle.HEADLINE_MEDIUM, color=ft.colors.PRIMARY),
            ft.Divider(color=ft.colors.PRIMARY),
            ft.Text("Мова розпізнавання:", style=ft.TextThemeStyle.TITLE_MEDIUM),
            self.language_dropdown,
            ft.Divider(color=ft.colors.PRIMARY),
            ft.Text("Тип розпізнавання:", style=ft.TextThemeStyle.TITLE_MEDIUM),
            self.recognition_type,
            ft.Divider(color=ft.colors.PRIMARY),
            ft.Text("Шляхи до Vosk моделей:", style=ft.TextThemeStyle.TITLE_MEDIUM),
            self.vosk_model_paths["uk-UA"],
            self.vosk_model_paths["en-US"],
            self.vosk_model_paths["ru-RU"],
            ft.Divider(color=ft.colors.PRIMARY),
            self.enable_tts,
            save_button,
        ], scroll=ft.ScrollMode.AUTO, expand=True)

        self.text_generation_model = ft.Dropdown(
            label="Модель генерації тексту",
            options=[
                ft.dropdown.Option("custom", "Кастомна модель"),
                ft.dropdown.Option("groq", "Системна модель"),
            ],
            value=self.config.get("text_generation_model", "custom"),
            on_change=self.change_text_generation_model,
            width=300,
            border_color=ft.colors.PRIMARY,
            focused_border_color=ft.colors.SECONDARY,
        )

        self.ai_settings = ft.Column([
            ft.Text("Налаштування штучного інтелекту", style=ft.TextThemeStyle.HEADLINE_MEDIUM, color=ft.colors.PRIMARY),
            ft.Divider(color=ft.colors.PRIMARY),
            self.use_ai,
            self.ai_model,
            self.text_generation_model,
            self.hf_api_key,
            self.groq_api_token,  # Add the new Groq API token field
            save_button,
        ], scroll=ft.ScrollMode.AUTO, expand=True)

        self.flux_settings = ft.Column([
            ft.Text("Налаштування моделі FLUX", style=ft.TextThemeStyle.HEADLINE_MEDIUM, color=ft.colors.PRIMARY),
            ft.Divider(color=ft.colors.PRIMARY),
            self.flux_model,
            save_button,
        ], scroll=ft.ScrollMode.AUTO, expand=True)

        self.settings_tabs = ft.CupertinoSlidingSegmentedButton(
            selected_index=0,
            thumb_color=ft.colors.PRIMARY,
            on_change=self.on_settings_tab_change,
            padding=ft.padding.symmetric(0, 10),
            controls=[
                ft.Text("Розпізнавання голосу", color=ft.colors.PRIMARY),
                ft.Text("ШІ", color=ft.colors.PRIMARY),
                ft.Text("FLUX", color=ft.colors.PRIMARY),
            ],
        )

        self.settings_content = ft.Container(
            content=self.voice_recognition_settings,
            expand=True,
            padding=20,
            border_radius=10,
            bgcolor=ft.colors.SURFACE_VARIANT,
        )

    def build(self):
        return ft.Column([
            self.settings_tabs,
            self.settings_content,
        ], expand=True)

    def on_settings_tab_change(self, e):
        if e.data == "0":
            self.settings_content.content = self.voice_recognition_settings
        elif e.data == "1":
            self.settings_content.content = self.ai_settings
        elif e.data == "2":
            self.settings_content.content = self.flux_settings
        self.page.update()

    def change_language(self, e):
        new_language = e.control.value
        self.voice_recognition.change_language(new_language)
        
        # Update the Vosk model path if using Vosk
        if self.voice_recognition.use_vosk:
            new_model_path = self.vosk_model_paths[new_language].value
            self.voice_recognition.set_vosk_model_path(new_language, new_model_path)
        
        # Update the config
        self.config["language"] = new_language
        save_config(self.config)
        
        # Restart listening to apply changes
        self.voice_recognition.restart_listening()

    def change_recognition_type(self, e):
        new_type = e.control.value
        self.voice_recognition.set_recognition_type(new_type)
        
        # If switching to Vosk, ensure the correct model is loaded
        if new_type == "vosk":
            current_language = self.language_dropdown.value
            new_model_path = self.vosk_model_paths[current_language].value
            self.voice_recognition.set_vosk_model_path(current_language, new_model_path)
        
        # Update config and restart listening
        self.config["preferred_recognition"] = new_type
        self.config["use_vosk"] = (new_type == "vosk")
        save_config(self.config)
        self.voice_recognition.restart_listening()


    def toggle_tts(self, e):
        config = load_config()
        config["enable_tts"] = e.control.value
        save_config(config)

    def toggle_ai(self, e):
        self.config["use_ai"] = e.control.value
        save_config(self.config)

    def change_text_generation_model(self, e):
        self.config["text_generation_model"] = e.control.value
        save_config(self.config)

    def save_settings(self, e):
        for lang, text_field in self.vosk_model_paths.items():
            self.voice_recognition.set_vosk_model_path(lang, text_field.value)
        
        self.config["ai_model"] = self.ai_model.value
        self.config["flux_model"] = self.flux_model.value
        self.config["text_generation_model"] = self.text_generation_model.value
        self.config["hf_api_key"] = self.hf_api_key.value
        self.config["groq_api_token"] = self.groq_api_token.value  # Save the Groq API token
        save_config(self.config)
        self.command_processor.reload_ai_config()
        self.page.show_snack_bar(ft.SnackBar(content=ft.Text("Налаштування збережено")))
        
        # Reload configuration
        self.config = load_config()
        self.voice_recognition.config = self.config

        self.voice_recognition.change_language(self.language_dropdown.value)
        self.voice_recognition.set_recognition_type(self.recognition_type.value)

        # Apply the new Groq API token immediately
        self.apply_groq_api_token()

        # Restart listening to apply all changes
        self.voice_recognition.restart_listening()

    def apply_groq_api_token(self):
        try:
            new_groq_client = Groq(api_key=self.config["groq_api_token"])
            # Here you would update the Groq client wherever it's used in your application
            # For example, if it's stored in the command_processor:
            self.command_processor.groq_client = new_groq_client
            logging.info("Groq API token updated and applied successfully.")
        except Exception as e:
            logging.error(f"Error applying new Groq API token: {str(e)}")
            self.page.show_snack_bar(ft.SnackBar(content=ft.Text("Помилка при оновленні Groq API токену")))

class ProfileView(BaseView):
    def __init__(self, page: ft.Page, username: str):
        super().__init__(page)
        self.username = username

    def build(self):
        return ft.Container(
            content=ft.Column([
                ft.Text(f"Профіль користувача: {self.username}", style=ft.TextThemeStyle.HEADLINE_MEDIUM, font_family="Roboto"),
                ft.Text("Планується", font_family="Roboto"),
            ]),
            alignment=ft.alignment.center,
            padding=20,
            expand=True,
        )

class ARTApp:
    def __init__(self, page: ft.Page):
        self.page = page
        self.main_view = MainView(page)
        self.weather_view = WeatherView(page)
        self.chat_view = ChatView(page)
        self.commands_view = CommandsApp(page)
        self.gestures_view = GesturesView(page)
        self.cards_view = CardsView(page)
        self.extras_view = ExtrasView(page)
        self.settings_view = SettingsView(page, self.main_view.voice_recognition, self.main_view.command_processor)
        self.profile_view = None  # Ініціалізуємо профіль як None
        self.content = None
        self.theme_icon = None
        self.menu = None
        self.header = None
        self.text_input = None
        self.send_button = None
        self.voice_input = None
        self.input_row = None
        self.current_view = None

        cards_view = CardsView(page)
        cards_view.loaded_cards = load_saved_cards()

    def build(self):
        self.page.title = "ART - Асистент розумних технологій"
        self.page.theme_mode = ft.ThemeMode.DARK
        self.page.window.width = 800
        self.page.window.height = 600
        self.page.padding = 0
        self.page.bgcolor = ft.colors.SURFACE_VARIANT
        self.page.window.on_close = self.on_window_close
        self.page.fonts = {
            "Roboto": "https://github.com/google/fonts/raw/main/apache/roboto/Roboto%5Bwdth,wght%5D.ttf"
        }

        self.theme_icon = ft.IconButton(
            icon=ft.icons.DARK_MODE if self.page.theme_mode == ft.ThemeMode.DARK else ft.icons.LIGHT_MODE,
            icon_color=ft.colors.PRIMARY,
            icon_size=30,
            on_click=self.toggle_theme,
            tooltip="Змінити тему",
        )

        self.menu = ft.PopupMenuButton(
            icon=ft.icons.MENU,
            items=[
                self.menu_item(ft.icons.HOME, "Головна"),
                self.menu_item(ft.icons.CLOUD, "Погода"),
                self.menu_item(ft.icons.CHAT, "Чат"),
                self.menu_item(ft.icons.LIST, "Команди"),
                self.menu_item(ft.icons.GESTURE, "Жести"),
                self.menu_item(ft.icons.CREDIT_CARD, "Картки"),
                # self.menu_item(ft.icons.ADD_CIRCLE, "Додатково"),
                self.menu_item(ft.icons.SETTINGS, "Налаштування"),
            ],
        )

        self.text_input = ft.TextField(
            hint_text="Введіть команду...",
            border_radius=30,
            expand=True,
            on_submit=lambda e: self.main_view.process_input(e.control.value),
            focused_border_color=ft.colors.PRIMARY,
            focused_bgcolor=ft.colors.PRIMARY_CONTAINER,
        )

        self.send_button = ft.IconButton(
            icon=ft.icons.SEND,
            icon_color=ft.colors.PRIMARY,
            icon_size=30,
            on_click=lambda _: self.main_view.process_input(self.text_input.value),
            tooltip="Надіслати",
        )

        self.voice_input = ft.IconButton(
            icon=ft.icons.MIC,
            icon_color=ft.colors.PRIMARY,
            icon_size=30,
            on_click=self.main_view.toggle_mic,
            tooltip="Голосове введення",
        )

        self.input_row = ft.Container(
            content=ft.Row(
                [self.text_input, self.send_button, self.voice_input],
                alignment=ft.MainAxisAlignment.CENTER,
            ),
            padding=20,
            bgcolor=ft.colors.SURFACE,
            shadow=ft.BoxShadow(blur_radius=5, color=ft.colors.BLACK12, offset=ft.Offset(0, -3)),
        )

        self.header = ft.Container(
            content=ft.Row(
                [
                    ft.Text("ART", style=ft.TextThemeStyle.HEADLINE_MEDIUM, color=ft.colors.PRIMARY, font_family="Roboto"),
                    self.menu,
                    ft.Container(expand=True),
                    self.theme_icon,
                    ft.IconButton(icon=ft.icons.ACCOUNT_CIRCLE, icon_color=ft.colors.PRIMARY, tooltip="Профіль", on_click=self.show_profile),  # Додаємо on_click для відображення профілю
                ],
                alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
            ),
            padding=15,
            bgcolor=ft.colors.SURFACE,
            shadow=ft.BoxShadow(blur_radius=5, color=ft.colors.BLACK12),
        )

        self.content = ft.Container(content=self.main_view.build(), expand=True)
        self.page.add(self.header, self.content)

    def on_window_close(self, e):
        print("Закриття застосунку...")
        self.voice_recognition.stop_listening()  # Зупиняємо прослуховування
        time.sleep(1)  # Додаємо невелику затримку для завершення потоків
        self.page.window.close()
        sys.exit()

    def toggle_theme(self, _):
        self.page.theme_mode = ft.ThemeMode.LIGHT if self.page.theme_mode == ft.ThemeMode.DARK else ft.ThemeMode.DARK
        self.theme_icon.icon = ft.icons.DARK_MODE if self.page.theme_mode == ft.ThemeMode.DARK else ft.icons.LIGHT_MODE
        self.main_view.update_message_colors()
        self.update_mic_state() 
        self.page.update()

    def menu_item(self, icon, text):
        return ft.PopupMenuItem(
            content=ft.Row([ft.Icon(icon, color=ft.colors.PRIMARY), ft.Text(text, font_family="Roboto")]),
            on_click=lambda _: self.switch_view(text.lower())
        )

    def switch_view(self, view_name):
        self.current_view = view_name

        if view_name == "головна":
            self.content.content = self.main_view.build()
        elif view_name == "погода":
            self.content.content = self.weather_view.build()
        elif view_name == "чат":
            self.content.content = self.chat_view.build()
        elif view_name == "команди":
            self.content.content = self.commands_view.build()
            self.commands_view.show_commands_list()
        elif view_name == "жести":
            self.content.content = self.gestures_view.build()
        elif view_name == "картки(пророблюється)":
            self.content.content = self.cards_view.build()
        # elif view_name == "додатково":
        #     self.content.content = self.extras_view.build()
        elif view_name == "налаштування":
            self.content.content = self.settings_view.build()
        else:
            self.content.content = self.main_view.build()

        self.page.update()
        self.update_mic_state()
        if view_name == "головна":
            self.main_view.update_message_colors()

    def update_mic_state(self):
        self.main_view.update_mic_state() 

    def start_permission_check(self):
        def check_permission_periodically():
            while True:
                time.sleep(600)  # Sleep for 10 minutes
                if not self.check_permission_and_notify():
                    break  # Exit the loop if permission is not granted

        thread = threading.Thread(target=check_permission_periodically, daemon=True)
        thread.start()

    def check_permission_and_notify(self):
        permission_checker = PermissionChecker()
        if not permission_checker.check_permission():
            self.show_permission_expired_dialog()
            return False
        return True

    def show_permission_expired_dialog(self):
        self.dialog_event = threading.Event()  # Event to block execution

        def on_dialog_close(_):
            self.dialog_event.set()  # Unblock execution when dialog is closed

        dialog = ft.AlertDialog(
            title=ft.Text("Доступ Закінчився"),
            content=ft.Text("Ваш доступ до застосунку закінчився. Будь ласка, зверніться до адміністратора."),
            actions=[
                ft.TextButton("OK", on_click=on_dialog_close)
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )
        self.page.dialog = dialog
        dialog.open = True
        self.page.update()

        self.dialog_event.wait()  # Block execution until the event is set
        self.page.window.close()  # Close the window after the dialog is closed

    def show_profile(self, e):
        if self.profile_view is None:
            self.profile_view = ProfileView(self.page, GLOBAL_USERNAME)  # Створюємо профіль з глобальним іменем користувача
        self.content.content = self.profile_view.build()
        self.page.update()

class PermissionChecker:
    def __init__(self):
        self.api_url = 'https://famous-discrete-lioness.ngrok-free.app'
        self.local_permission_file = 'permission.pkl'
        self.credentials_file = 'credentials.json'
        self.permission_data = self.load_local_permission()
        self.credentials = self.load_credentials()
        self.load_global_credentials()

    def get_system_id(self):
        try:
            system_id = subprocess.check_output("wmic bios get serialnumber", shell=True)
            system_id = system_id.decode().split("\n")[1].strip()
            return system_id if system_id else "Unknown ID"
        except Exception as e:
            return f"Error: {str(e)}"

    def load_local_permission(self):
        if os.path.exists(self.local_permission_file):
            try:
                with open(self.local_permission_file, 'rb') as f:
                    return pickle.load(f)
            except (pickle.UnpicklingError, IOError):
                pass
        return None

    def save_local_permission(self, permission_data):
        try:
            with open(self.local_permission_file, 'wb') as f:
                pickle.dump(permission_data, f)
        except IOError:
            pass

    def load_credentials(self):
        if os.path.exists(self.credentials_file):
            try:
                with open(self.credentials_file, 'r') as f:
                    return json.load(f)
            except json.JSONDecodeError:
                pass
        return {}

    def save_credentials(self, username, password):
        global GLOBAL_USERNAME, GLOBAL_PASSWORD
        credentials = {
            'username': username,
            'password': password
        }
        try:
            with open(self.credentials_file, 'w') as f:
                json.dump(credentials, f)
            GLOBAL_USERNAME = username
            GLOBAL_PASSWORD = password
            print(f"Збережено нові облікові дані. Користувач: {GLOBAL_USERNAME}")
        except IOError:
            print("Помилка при збереженні облікових даних")

    def load_global_credentials(self):
        global GLOBAL_USERNAME, GLOBAL_PASSWORD
        if self.credentials:
            GLOBAL_USERNAME = self.credentials.get('username')
            GLOBAL_PASSWORD = self.credentials.get('password')
            print(f"Завантажено збережені облікові дані. Користувач: {GLOBAL_USERNAME}")

    def check_permission(self):
        global GLOBAL_USERNAME, GLOBAL_PASSWORD
        if not self.credentials:
            return False
        username = self.credentials.get('username')
        password = self.credentials.get('password')
        device_id = self.get_system_id()

        print(f"Перевірка доступу для користувача: {username}")

        try:
            response = requests.post(f'{self.api_url}/check_permission', 
                                    json={'device_id': device_id, 'username': username, 'password': password},
                                    timeout=5)  # Додаємо таймаут
            response.raise_for_status()
            data = response.json()

            if data.get('permission', False):
                self.permission_data = data
                self.save_local_permission(data)
                self.save_credentials(username, password)
                return True
            return False
        except (requests.RequestException, requests.Timeout) as e:
            print(f"Помилка при перевірці доступу онлайн: {str(e)}. Перевірка локальних даних.")
            return self.check_local_permission()

    def check_local_permission(self):
        if self.permission_data:
            if 'expiration_date' in self.permission_data:
                try:
                    expiration_date = datetime.strptime(self.permission_data['expiration_date'].split('T')[0], '%Y-%m-%d')
                    if expiration_date > datetime.now():
                        print("Локальний дозвіл дійсний.")
                        return True
                    else:
                        print("Локальний дозвіл застарів.")
                except ValueError as e:
                    print(f"Помилка при обробці дати: {str(e)}")
            else:
                print("Дата закінчення відсутня в локальних даних. Надаємо тимчасовий доступ.")
                return True
        print("Локальні дані про дозвіл відсутні.")
        return False


    def register_user(self, username, password):
        device_id = self.get_system_id()
        try:
            response = requests.post(f'{self.api_url}/register', 
                                     json={'device_id': device_id, 'username': username, 'password': password})
            response.raise_for_status()
            data = response.json()
            return data.get('success', False), data.get('message', '')
        except requests.RequestException as e:
            return False, str(e)
        
speak_text = CommandProcessor(None)
speak_text.speak("Вітаємо вас у застосунку АРТ!")
# speak_text.speak("Для початку діалогу натисніть на мікрофон або введіть команду вручну.")

def main(page: ft.Page):
    page.title = "Розумний Помічник"
    page.theme_mode = ft.ThemeMode.LIGHT
    page.window.width = 400
    page.window.height = 630
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER
    page.bgcolor = ft.colors.INDIGO_50

    permission_checker = PermissionChecker()

    def launch_art_app():
        page.controls.clear()
        art_app = ARTApp(page)
        art_app.build()
        art_app.start_permission_check()
        
        # Налаштовуємо обробник закриття вікна
        page.window.on_close = art_app.on_window_close
        
        page.update()

    def show_login_form():
        def on_login(e):
            login_button.scale = ft.Scale(0.9)
            page.update()

            permission_checker.credentials = {
                'username': username_field.value,
                'password': password_field.value
            }

            if permission_checker.check_permission():
                print("Доступ надано. Перехід до ARTApp.")
                permission_checker.save_credentials(username_field.value, password_field.value)
                launch_art_app()
            else:
                print("Доступ заборонено.")
                snack = ft.SnackBar(content=ft.Text("Доступ заборонено. Перевірте дані або зареєструйтесь. Можливо у вас закінчився доступ."))
                page.overlay.append(snack)
                snack.open = True
                page.update()

            login_button.scale = ft.Scale(1.0)
            page.update()

        def on_register(e):
            if len(password_field.value) < 6:
                snack = ft.SnackBar(content=ft.Text("Пароль повинен містити мінімум 6 символів."))
                page.overlay.append(snack)
                snack.open = True
                page.update()
                return

            success, message = permission_checker.register_user(username_field.value, password_field.value)
            if success:
                snack = ft.SnackBar(content=ft.Text("Реєстрація успішна. Тепер ви можете увійти."))
            else:
                snack = ft.SnackBar(content=ft.Text(f"Помилка реєстрації: {message}"))
            page.overlay.append(snack)
            snack.open = True
            page.update()

        def styled_textfield(label, password=False):
            return ft.TextField(
                label=label,
                border_radius=12,
                width=300,
                password=password,
                border_color=ft.colors.INDIGO_400,
                focused_border_color=ft.colors.INDIGO_600,
                prefix_icon=ft.icons.PERSON if not password else ft.icons.LOCK,
            )

        username_field = styled_textfield("Ім'я користувача")
        password_field = styled_textfield("Пароль", password=True)

        # Fill in saved credentials if available
        if permission_checker.credentials:
            username_field.value = permission_checker.credentials.get('username', '')
            password_field.value = permission_checker.credentials.get('password', '')

        computer_id = permission_checker.get_system_id()

        def copy_id(e):
            page.set_clipboard(computer_id)
            snack = ft.SnackBar(content=ft.Text("ID скопійовано!"))
            page.overlay.append(snack)
            snack.open = True
            page.update()

        tooltip_text = "Для того щоб увійти у застосунок передайте код адміністратору (телеграм: @kiril365) для надання доступу до застосунку"

        id_container = ft.Container(
            content=ft.Row(
                [
                    ft.Text(computer_id, size=12, color=ft.colors.INDIGO_400),
                    ft.IconButton(
                        icon=ft.icons.COPY,
                        icon_color=ft.colors.INDIGO_400,
                        icon_size=20,
                        on_click=copy_id,
                        tooltip="Копіювати ID",
                    ),
                ],
                alignment=ft.MainAxisAlignment.CENTER,
            ),
            tooltip=tooltip_text,
            padding=10,
            border_radius=8,
            ink=True,
        )

        username_container = ft.Container(
            content=username_field,
            padding=10,
            scale=ft.Scale(0.8),
            animate_scale=ft.Animation(600, ft.AnimationCurve.ELASTIC_OUT),
        )
        password_container = ft.Container(
            content=password_field,
            padding=10,
            scale=ft.Scale(0.8),
            animate_scale=ft.Animation(600, ft.AnimationCurve.ELASTIC_OUT),
        )

        def animate_containers(e):
            username_container.scale = ft.Scale(1.0)
            password_container.scale = ft.Scale(1.0)
            page.update()

        login_button = ft.ElevatedButton(
            text="Увійти",
            on_click=on_login,
            width=200,
            height=30,
            style=ft.ButtonStyle(
                color=ft.colors.WHITE,
                bgcolor=ft.colors.INDIGO_600,
                shape=ft.RoundedRectangleBorder(radius=12),
            ),
            animate_scale=ft.Animation(200, ft.AnimationCurve.EASE_IN_OUT),
        )

        logo = ft.Icon(ft.icons.LOCK, size=80, color=ft.colors.INDIGO_600)

        form_container = ft.Container(
            content=ft.Column(
                [
                    logo,
                    ft.Text("Вітаємо!", size=24, weight=ft.FontWeight.BOLD, color=ft.colors.INDIGO_600),
                    ft.Text("Будь ласка, увійдіть до свого облікового запису", size=14, color=ft.colors.INDIGO_400),
                    username_container,
                    password_container,
                    id_container,
                    login_button,
                ],
                horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                spacing=20,
            ),
            width=350,
            height=550,
            padding=30,
            border_radius=20,
            alignment=ft.alignment.center,
            bgcolor=ft.colors.WHITE,
            shadow=ft.BoxShadow(spread_radius=1, blur_radius=15, color=ft.colors.with_opacity(0.2, ft.colors.INDIGO)),
            animate=ft.Animation(1000, ft.AnimationCurve.ELASTIC_OUT),
        )
        register_button = ft.ElevatedButton(
            text="Зареєструватися",
            on_click=on_register,
            width=200,
            height=30,
            style=ft.ButtonStyle(
                color=ft.colors.WHITE,
                bgcolor=ft.colors.INDIGO_400,
                shape=ft.RoundedRectangleBorder(radius=12),
            ),
            animate_scale=ft.Animation(200, ft.AnimationCurve.EASE_IN_OUT),
        )

        form_container = ft.Container(
            content=ft.Column(
                [
                    logo,
                    ft.Text("Вітаємо!", size=24, weight=ft.FontWeight.BOLD, color=ft.colors.INDIGO_600),
                    ft.Text("Будь ласка, увійдіть або зареєструйтесь", size=14, color=ft.colors.INDIGO_400),
                    username_container,
                    password_container,
                    id_container,
                    login_button,
                    register_button,
                ],
                horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                spacing=20,
            ),
            width=350,
            height=600,
            padding=30,
            border_radius=20,
            alignment=ft.alignment.center,
            bgcolor=ft.colors.WHITE,
            shadow=ft.BoxShadow(spread_radius=1, blur_radius=15, color=ft.colors.with_opacity(0.2, ft.colors.INDIGO)),
            animate=ft.Animation(1000, ft.AnimationCurve.ELASTIC_OUT),
        )

        page.add(form_container)
        page.on_load = animate_containers


    if permission_checker.check_permission():
        print("Доступ надано. Запуск ARTApp.")
        launch_art_app()
    else:
        print("Доступ не надано. Показ форми логіну.")
        show_login_form()

if __name__ == "__main__":
    ft.app(target=main)